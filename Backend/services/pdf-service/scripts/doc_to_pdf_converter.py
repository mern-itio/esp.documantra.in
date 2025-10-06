#!/usr/bin/env python3
"""
DOC to PDF Converter using python-docx and reportlab
This script provides high-fidelity DOC to PDF conversion with layout preservation
"""

import sys
import os
import argparse
from pathlib import Path

def convert_doc_to_pdf(input_path, output_path):
    """
    Convert DOC/DOCX to PDF using LibreOffice for better image and layout preservation
    
    Args:
        input_path (str): Path to input DOC/DOCX file
        output_path (str): Path to output PDF file
    
    Returns:
        dict: Conversion result with success status and details
    """
    try:
        import subprocess
        import tempfile
        import shutil
        
        # Check if input file exists
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input DOC/DOCX file not found: {input_path}")
        
        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Use LibreOffice for conversion (better image and layout preservation)
        temp_dir = tempfile.mkdtemp()
        try:
            # Enhanced LibreOffice command with headless VCL and no GUI flags
            cmd = [
                'libreoffice',
                '--headless',
                '--nodefault',
                '--nolockcheck',
                '--nologo',
                '--norestore',
                '--nofirststartwizard',
                '-env:UserInstallation=file:///tmp/lo_profile_py',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', temp_dir,
                input_path
            ]
            
            # Set environment variables for better LibreOffice conversion
            env = os.environ.copy()
            env.update({
                'HOME': '/tmp',
                'DISPLAY': '',
                'SAL_USE_VCLPLUGIN': 'headless',
                'SAL_DISABLE_OPENCL': '1',
                'SAL_DISABLE_OPENCL_IMAGING': '1'
            })
            
            # Execute LibreOffice conversion
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, env=env)
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

            # Fallback: if no output and xvfb-run exists, try it
            out_file = os.path.splitext(os.path.basename(input_path))[0] + '.pdf'
            temp_pdf_path = os.path.join(temp_dir, out_file)
            if not os.path.exists(temp_pdf_path):
                try:
                    _ = subprocess.run(['which', 'xvfb-run'], capture_output=True, text=True)
                    xvfb_cmd = ['xvfb-run', '-a'] + cmd
                    result2 = subprocess.run(xvfb_cmd, capture_output=True, text=True, timeout=300, env=env)
                    if result2.returncode != 0:
                        raise RuntimeError(f"xvfb-run LibreOffice conversion failed: {result2.stderr}")
                except Exception as _e:
                    pass
            
            # Find the generated PDF file
            input_filename = os.path.splitext(os.path.basename(input_path))[0]
            temp_pdf_path = os.path.join(temp_dir, f"{input_filename}.pdf")
            
            if not os.path.exists(temp_pdf_path):
                raise RuntimeError("LibreOffice did not generate output PDF file")
            
            # Move the generated PDF to the desired output location
            shutil.move(temp_pdf_path, output_path)
            
            # Get file sizes
            input_size = os.path.getsize(input_path)
            output_size = os.path.getsize(output_path)
            
            return {
                "success": True,
                "input_size": input_size,
                "output_size": output_size,
                "message": "DOC/DOCX converted to PDF successfully using LibreOffice with full layout and image preservation",
                "conversion_method": "LibreOffice CLI (Python wrapper)"
            }
            
        finally:
            # Clean up temporary directory
            shutil.rmtree(temp_dir, ignore_errors=True)
        
    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "error": "LibreOffice conversion timed out",
            "message": "Conversion timed out after 5 minutes"
        }
    except FileNotFoundError:
        # LibreOffice not found, try fallback
        print("LibreOffice not found, falling back to basic conversion")
        return convert_doc_to_pdf_fallback(input_path, output_path)
    except Exception as e:
        print(f"Error in LibreOffice conversion: {e}")
        # Try fallback method
        return convert_doc_to_pdf_fallback(input_path, output_path)

def convert_doc_to_pdf_fallback(input_path, output_path):
    """
    Fallback DOC to PDF conversion using python-docx and reportlab with image support
    
    Args:
        input_path (str): Path to input DOC/DOCX file
        output_path (str): Path to output PDF file
    
    Returns:
        dict: Conversion result with success status and details
    """
    try:
        # Import required libraries
        from docx import Document
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        import tempfile
        
        # Read DOCX file
        doc = Document(input_path)
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=8,
            spaceBefore=12
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_JUSTIFY
        )
        
        # Process paragraphs and images
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue
            
            # Determine style based on paragraph properties
            if paragraph.style.name.startswith('Title'):
                story.append(Paragraph(text, title_style))
            elif paragraph.style.name.startswith('Heading'):
                story.append(Paragraph(text, heading_style))
            else:
                story.append(Paragraph(text, normal_style))
        
        # Process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(Spacer(1, 12))
                story.append(pdf_table)
                story.append(Spacer(1, 12))
        
        # Process images (extract and embed)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    if image_data:
                        # Save image to temporary file
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                            temp_file.write(image_data)
                            temp_image_path = temp_file.name
                        
                        # Add image to PDF
                        img = Image(temp_image_path, width=4*inch, height=3*inch)
                        story.append(Spacer(1, 12))
                        story.append(img)
                        story.append(Spacer(1, 12))
                        
                        # Clean up temporary file
                        os.unlink(temp_image_path)
                except Exception as e:
                    print(f"Warning: Could not process image: {e}")
                    continue
        
        # Build PDF
        pdf_doc.build(story)
        
        # Get file sizes
        input_size = os.path.getsize(input_path)
        output_size = os.path.getsize(output_path)
        
        return {
            "success": True,
            "input_size": input_size,
            "output_size": output_size,
            "message": "DOC/DOCX converted to PDF using fallback method with image support",
            "conversion_method": "Fallback with image extraction",
            "warning": "Layout preservation may be limited with fallback method"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": "Both primary and fallback conversion methods failed"
        }

def main():
    """Main function to handle command line arguments and execute conversion"""
    parser = argparse.ArgumentParser(description='Convert DOC/DOCX to PDF')
    parser.add_argument('input_path', help='Path to input DOC/DOCX file')
    parser.add_argument('output_path', help='Path to output PDF file')
    
    args = parser.parse_args()
    
    # Convert DOC to PDF
    result = convert_doc_to_pdf(args.input_path, args.output_path)
    
    # Print result as JSON for Node.js to parse
    import json
    print(json.dumps(result, indent=2))
    
    # Exit with appropriate code
    sys.exit(0 if result.get('success', False) else 1)

if __name__ == "__main__":
    main()

