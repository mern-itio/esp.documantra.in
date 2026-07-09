#!/usr/bin/env python3
"""Generate Annexure A4 — Information Security Policy (DocuMantra ASP)."""
from datetime import datetime, timezone
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "deploy" / "docs" / "asp-audit-annexures"
OUT = OUT_DIR / "Annexure-A4-Information-Security-Policy.docx"
OUT_ALIAS = OUT_DIR / "Information-Security-Policy.docx"

ORG = "ITIO Innovex Pvt Ltd"
PRODUCT = "DocuMantra (ESP)"
URL = "https://esp.documantra.in"
VERSION = "1.1"
DOCUMENT_ID = "ISP-DOCUMANTRA-A4"
EFFECTIVE = "09 July 2026"
REVIEW = "09 July 2027"
APPROVAL_DATE = "09 July 2026"
REGISTERED = (
    "C-32, Sector-14, Kaushambi, Ghaziabad, Uttar Pradesh 201010, India "
    "(CIN: U72300UP2015PTC070364)"
)

CLASSIFICATION_LABEL = "CONFIDENTIAL"
CLASSIFICATION_DETAIL = (
    "Internal / Auditor shareable — redact secrets, credentials, and customer PII "
    "before distribution outside ITIO Innovex Pvt Ltd or accredited auditors."
)

PREPARED_BY = "Information Security Owner / Technical Lead"
REVIEWED_BY = "Information Security Owner"
APPROVED_BY = "Management Representative"
APPROVED_BY_TITLE = "Director / Authorized Signatory, ITIO Innovex Pvt Ltd"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_control_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, value) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value


def build() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    class_banner = doc.add_paragraph()
    class_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    class_run = class_banner.add_run(CLASSIFICATION_LABEL)
    class_run.bold = True
    class_run.font.size = Pt(14)
    class_run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    class_detail = doc.add_paragraph()
    class_detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail_run = class_detail.add_run(CLASSIFICATION_DETAIL)
    detail_run.font.size = Pt(9)
    detail_run.italic = True

    doc.add_paragraph()

    title = doc.add_heading("Information Security Policy", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{ORG}\n{PRODUCT} — Application Service Provider (ASP)\n")
    run.bold = True
    sub.add_run(f"Production URL: {URL}\n")
    sub.add_run(
        f"Document ID: {DOCUMENT_ID} | Version: {VERSION} | "
        f"Effective: {EFFECTIVE} | Next review: {REVIEW}"
    )

    doc.add_paragraph()

    add_heading(doc, "Document control", 1)
    add_control_table(
        doc,
        [
            ("Document title", "Information Security Policy"),
            ("Document ID", DOCUMENT_ID),
            ("Annexure reference", "A4 — Information-Security-Policy.docx"),
            ("Version", VERSION),
            ("Effective date", EFFECTIVE),
            ("Next review date", REVIEW),
            ("Organization", ORG),
            ("Registered address", REGISTERED),
            ("Document owner", PREPARED_BY),
            ("Classification", f"{CLASSIFICATION_LABEL} — {CLASSIFICATION_DETAIL}"),
            ("Distribution", "Management, operations, development, and ASP/CERT-In auditors under NDA"),
            ("Approval authority", APPROVED_BY_TITLE),
            ("Approval date", APPROVAL_DATE),
        ],
    )

    doc.add_paragraph()

    add_heading(doc, "Document classification scheme", 1)
    doc.add_paragraph(
        "All information security and system documentation for DocuMantra is labelled using "
        "the following classification levels. Handlers must apply the highest applicable level."
    )
    add_control_table(
        doc,
        [
            ("Public", "Approved for unrestricted release (marketing, public website content)."),
            (
                "Internal",
                "For employees and contractors of ITIO Innovex; not for public distribution.",
            ),
            (
                "Confidential",
                "Sensitive business, security architecture, or audit evidence; share with auditors "
                "only under NDA after redaction of secrets.",
            ),
            (
                "Restricted",
                "Cryptographic keys, production credentials, live customer PII, or ESP tokens; "
                "never included in annexures or email.",
            ),
        ],
    )
    doc.add_paragraph(
        f"This policy document is classified as: {CLASSIFICATION_LABEL} ({CLASSIFICATION_DETAIL})"
    )

    doc.add_paragraph()

    add_heading(doc, "1. Purpose", 1)
    doc.add_paragraph(
        "This Information Security Policy establishes the security objectives, principles, and "
        "minimum controls for the DocuMantra electronic signature and document workflow platform "
        "operated by ITIO Innovex Pvt Ltd as an Application Service Provider (ASP). The policy "
        "supports protection of customer data, integrity of electronic signatures, and compliance "
        "with applicable Indian laws and contractual obligations with ESPs and customers."
    )

    add_heading(doc, "2. Scope", 1)
    add_bullets(
        doc,
        [
            f"Production environment: {URL} and associated APIs, admin console, and public signer flows.",
            "All personnel, contractors, and third parties with access to production systems, source code, or customer data.",
            "Microservices architecture: authentication, e-sign, document, identity, subscription, and admin services.",
            "Integrations with ESP (eSign), DigiLocker/KYC providers, and cloud infrastructure (DigitalOcean).",
            "Exclusions: end-user devices and networks outside the organization's control; customer-owned content policy.",
        ],
    )

    add_heading(doc, "3. Roles and responsibilities", 1)
    add_bullets(
        doc,
        [
            "Management: approves policy, allocates resources, and accepts residual risk.",
            "Information Security Owner / Technical Lead: implements controls, coordinates VAPT and incident response.",
            "Development team: secure SDLC, code review, dependency management, and remediation of findings.",
            "Operations / DevOps: hardening, patching, backups, monitoring, and access to production.",
            "All users: comply with password, 2FA, and acceptable-use requirements.",
        ],
    )

    add_heading(doc, "4. Information security principles", 1)
    add_numbered(
        doc,
        [
            "Confidentiality — access only on need-to-know basis; encryption in transit.",
            "Integrity — tamper-evident signing, audit trails, and input validation.",
            "Availability — resilient hosting, monitoring, and recovery procedures.",
            "Accountability — authentication, logging, and traceability of envelope actions.",
            "Compliance — alignment with IT Act 2000, IT Rules 2011, and CERT-In directions where applicable.",
        ],
    )

    add_heading(doc, "5. Asset management", 1)
    add_bullets(
        doc,
        [
            "Application assets include source code (Git), Docker images, configuration, TLS certificates, and database records.",
            "Customer documents and signature evidence are stored in controlled storage with access restricted to authorized services.",
            "Cryptographic keys and API tokens are stored in environment variables or secure vaults; not committed to source control.",
            "Asset inventory maintained via repository structure, docker-compose service definitions, and deployment documentation.",
        ],
    )

    add_heading(doc, "6. Access control", 1)
    add_bullets(
        doc,
        [
            "Role-based access for application features (user, admin, organization roles).",
            "Production server access limited to authorized administrators; SSH key-based authentication.",
            "Database (MongoDB) not exposed to the public internet; bound to internal network interfaces.",
            "Microservice ports bound to localhost; external access only through nginx reverse proxy.",
            "Principle of least privilege for service accounts and API keys.",
            "Maximum concurrent sessions per user account: configurable (default 5) per organizational requirement.",
        ],
    )

    add_heading(doc, "7. Authentication and identity", 1)
    add_bullets(
        doc,
        [
            "User authentication via email/password with JWT bearer tokens for API access.",
            "Password policy: minimum length, complexity (upper, lower, number, special), and password history to prevent reuse.",
            "Two-factor authentication (TOTP): available for users and administrators; grace period for existing accounts before enforcement prompt.",
            "Admin panel: separate authentication surface with stricter session controls.",
            "Account lockout and rate limiting on authentication endpoints to mitigate brute-force attacks.",
            "Session idle timeout: configurable via admin session policy (default 8 hours).",
        ],
    )

    add_heading(doc, "8. Cryptography and transport security", 1)
    add_bullets(
        doc,
        [
            f"All public access over HTTPS/TLS 1.2+ at {URL}; HTTP redirected to HTTPS.",
            "HSTS enabled on user and admin applications.",
            "Electronic signatures use ESP-standard APIs (eSign 2.1); ASP signs request XML with RSA-SHA256 XML-DSig.",
            "Signed documents use PKCS7 signatures from ESP; document hashes use SHA-256.",
            "Secrets (JWT signing keys, API bearer tokens, DB credentials) rotated on compromise or role change.",
        ],
    )

    add_heading(doc, "9. Application security", 1)
    add_bullets(
        doc,
        [
            "Secure development practices: input validation, output encoding, SSRF guards on URL ingestion (e.g. organization logo).",
            "Security headers: Content-Security-Policy, X-Frame-Options, X-Content-Type-Options on SPAs.",
            "File upload validation with type/size presets for e-sign documents.",
            "CORS and API authentication enforced on protected routes; public signer routes scoped to envelope tokens.",
            "Webhook endpoints validated; DigiLocker callback routes monitored.",
            "Dependency vulnerability management via npm audit and periodic updates.",
            "VAPT (Web Application Security Testing) conducted June 2026; all 32 findings closed in final retest (v1.1).",
        ],
    )

    add_heading(doc, "10. Logging, monitoring, audit, and consent", 1)
    add_bullets(
        doc,
        [
            "Envelope and signing actions recorded in MongoDB audittrails collection.",
            "User consent records stored in userconsents collection (Terms of Service, Privacy Policy, marketing opt-in, e-sign disclosure) with IP address, user agent, version, and timestamp.",
            "Authentication events and security-sensitive operations logged server-side.",
            "Server and application logs retained per operational policy; CERT-In log retention (180 days) recommended for production.",
            "Periodic review of failed login attempts, 4xx/5xx rates, and infrastructure alerts.",
        ],
    )

    add_heading(doc, "11. Network and infrastructure security", 1)
    add_bullets(
        doc,
        [
            "Hosting: DigitalOcean cloud (production droplet); nginx reverse proxy terminates TLS.",
            "Firewall: cloud firewall and UFW — inbound ports limited (typically 22, 80, 443).",
            "Docker containerization for application services; non-root execution where practicable.",
            "Automatic security updates (unattended-upgrades) on Ubuntu host.",
            "fail2ban or equivalent for SSH brute-force mitigation.",
            "Linux cloud endpoint protection model: defence-in-depth controls in lieu of desktop antivirus agents on API servers (see Annexure A6/A15).",
            "No direct exposure of internal APIs or database ports to the internet.",
        ],
    )

    add_heading(doc, "12. Third-party and ESP integration", 1)
    add_bullets(
        doc,
        [
            "ESP integration via standard eSign APIs (XML request/response per CCA specification).",
            "KYC/DigiLocker integration via authorized API providers (HTTPS, bearer token authentication).",
            "Due diligence on subprocessors; data processing limited to transaction necessity.",
            "API keys stored securely; sandbox and production tokens segregated.",
            "Contracts and agreements with ESPs maintained (see Annexure A14 index).",
        ],
    )

    add_heading(doc, "13. Data protection and privacy", 1)
    add_bullets(
        doc,
        [
            "Personal data processed only for authentication, signing, and compliance purposes.",
            "Aadhaar/eKYC data handled per ESP and UIDAI applicable guidelines; minimal retention.",
            "Customer documents accessible only to authorized envelope participants and administrators per role.",
            "Data deletion and retention aligned with customer agreements and legal requirements.",
            "Cross-border hosting (US datacenter) disclosed to customers; appropriate safeguards applied.",
        ],
    )

    add_heading(doc, "14. Incident management", 1)
    add_bullets(
        doc,
        [
            "Security incidents reported to Information Security Owner without undue delay.",
            "Containment: disable compromised accounts, rotate credentials, block malicious IPs.",
            "Investigation using logs, audit trails, and VAPT evidence where applicable.",
            "Customer notification when personal data or signing integrity is affected, per legal obligation.",
            "Post-incident review and control improvements documented.",
        ],
    )

    add_heading(doc, "15. Business continuity and backup", 1)
    add_bullets(
        doc,
        [
            "Infrastructure snapshots and database backups per operational schedule.",
            "Documented deployment and rollback procedures (Docker, nginx, branch-based releases).",
            "Recovery time objectives defined operationally; critical path is e-sign availability for active envelopes.",
        ],
    )

    add_heading(doc, "16. Compliance and audit", 1)
    add_bullets(
        doc,
        [
            "ASP audit requirements: preliminary information, annexures A1–A15, and ESP integration evidence.",
            "OWASP-aligned secure coding; annual or change-triggered VAPT recommended.",
            "Policy exceptions documented and approved by management (e.g. deferred React Router upgrade per VAPT business acceptance).",
        ],
    )

    add_heading(doc, "17. Policy review", 1)
    doc.add_paragraph(
        f"This policy is reviewed at least annually or upon significant architectural, regulatory, "
        f"or incident-driven changes. Next scheduled review: {REVIEW}. "
        f"Version changes require review by the Information Security Owner and approval by Management."
    )

    add_heading(doc, "18. Approval and authorization", 1)
    doc.add_paragraph(
        "The undersigned confirm that this Information Security Policy has been reviewed, "
        "classified as documented above, and approved for implementation."
    )

    approval_table = doc.add_table(rows=4, cols=4)
    approval_table.style = "Table Grid"
    headers = ["Role", "Name", "Signature", "Date"]
    for col, header in enumerate(headers):
        approval_table.rows[0].cells[col].text = header

    approval_rows = [
        ("Prepared by", PREPARED_BY, "", EFFECTIVE),
        ("Reviewed by", REVIEWED_BY, "", EFFECTIVE),
        ("Approved by", f"{APPROVED_BY}\n{APPROVED_BY_TITLE}", "", APPROVAL_DATE),
    ]
    for row_idx, (role, name, signature, date) in enumerate(approval_rows, start=1):
        approval_table.rows[row_idx].cells[0].text = role
        approval_table.rows[row_idx].cells[1].text = name
        approval_table.rows[row_idx].cells[2].text = signature
        approval_table.rows[row_idx].cells[3].text = date

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: Physical or digital signature may be applied to the signed PDF copy submitted "
        "to auditors. This controlled document remains classified "
        f"{CLASSIFICATION_LABEL} until reclassified by Management."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ft = footer.add_run(
        f"\nGenerated {datetime.now(timezone.utc).strftime('%d %B %Y')} | "
        f"{ORG} | {PRODUCT} | {DOCUMENT_ID} v{VERSION} | {CLASSIFICATION_LABEL}"
    )
    ft.font.size = Pt(9)

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build()
    document.save(OUT)
    shutil.copy2(OUT, OUT_ALIAS)
    print(f"Written: {OUT}")
    print(f"Copied:  {OUT_ALIAS}")


if __name__ == "__main__":
    main()
