#!/usr/bin/env python3
"""Generate Annexure A7 — ASP-ESP eSign integration evidence (DOCX index)."""
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
ANNEX = ROOT / "deploy" / "docs" / "asp-audit-annexures"
OUT = ANNEX / "Annexure-A7-ASP-ESP-eSign-Integration.docx"

ORG = "ITIO Innovex Pvt Ltd"
PRODUCT = "Documantra (ASP)"
URL = "https://esp.documantra.in"


def main() -> None:
    doc = Document()
    h = doc.add_heading("Annexure A7 — ASP–ESP Integration via Standard eSign APIs", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{ORG} | {PRODUCT} | {URL}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y')}")
    doc.add_paragraph()

    doc.add_heading("Statement for ASP audit", 1)
    doc.add_paragraph(
        "Integration between the ASP (Documantra) and ESPs (e.g. V-Sign / NSDL eSign) is "
        "implemented via standard CCA eSign API version 2.1. The ASP initiates transactions "
        "using the ESP utility API (gettxnrefv4_1), signs the eSign request XML with ASP "
        "credentials (XML-DSig, RSA-SHA256), and receives the ESP response XML (EsignResp) "
        "containing PKCS7 document signatures on the configured callback URL."
    )

    doc.add_heading("Integration flow", 1)
    for step in [
        "User initiates Aadhaar-based signature on an envelope document.",
        "ASP computes document hash (SHA-256) and builds standard <Esign> request XML.",
        "ASP signs request XML using ASP private key (signXML.js — enveloped XML-DSig).",
        "ASP calls ESP utility: POST {UTILITY_URL}/gettxnrefv4_1 with aspId, txn, responseUrl, pdfdetails.",
        "Signer completes Aadhaar OTP on ESP auth page.",
        "ESP posts EsignResp XML to ASP callback: POST /esign/api/e-sign/public/v-sign/response.",
        "ASP parses txn, appends PKCS7 signature to PDF, updates audit trail.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Technical parameters", 1)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    rows = [
        ("API version", "eSign 2.1 (ver=21 in utility payload)"),
        ("ASP ID", "IIPLUAT001 (configurable via ASP_ID env)"),
        ("Auth mode", "1 — Aadhaar OTP"),
        ("Request root element", "<Esign> with <Docs><InputHash>…</InputHash></Docs>"),
        ("Request signature", "XML-DSig enveloped, RSA-SHA256, SHA-256 digest"),
        ("Response root element", "<EsignResp> with <Signatures><DocSignature sigType=PKCS7>"),
        ("Callback URL", f"{URL}/esign/api/e-sign/public/v-sign/response"),
        ("Code references", "helpers/signXML.js, services/signing/aadhar.vsign.service.js, controllers/vSignController.js"),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_heading("Sample XML files (attached in this annexure folder)", 1)
    doc.add_paragraph("Request (ASP → ESP): Annexure-A7-Sample-EsignReq.xml", style="List Bullet")
    doc.add_paragraph("Response (ESP → ASP): Annexure-A7-Sample-EsignResp.xml", style="List Bullet")
    doc.add_paragraph("Transaction init JSON (utility): Annexure-A8-ESP-gettxnref-request.json", style="List Bullet")

    doc.add_heading("Sample Request XML (redacted)", 1)
    req = (ANNEX / "Annexure-A7-Sample-EsignReq.xml").read_text(encoding="utf-8")
    p = doc.add_paragraph()
    run = p.add_run(req)
    run.font.name = "Consolas"
    run.font.size = run.font.size

    doc.add_page_break()
    doc.add_heading("Sample Response XML (redacted)", 1)
    resp = (ANNEX / "Annexure-A7-Sample-EsignResp.xml").read_text(encoding="utf-8")
    p2 = doc.add_paragraph()
    run2 = p2.add_run(resp)
    run2.font.name = "Consolas"

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
