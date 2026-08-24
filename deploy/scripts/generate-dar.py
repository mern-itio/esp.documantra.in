#!/usr/bin/env python3
"""Generate Daily Activity Report (DAR) for DocuMantra / ESP project."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "deploy" / "docs"

ORG = "ITIO Innovex Pvt Ltd"
PRODUCT = "DocuMantra (ESP)"
URL = "https://esp.documantra.in"
BRANCH = "feature/aadhaar-dual-appearance-vsign-live-21-8-26"
BASE_BRANCH = "recipient-portal-pandadoc-ux-10-7-26"
PREPARED_BY = "Technical Lead / Development Team"
REPORT_DATE = "21 August 2026"
OUT = OUT_DIR / "DAR-21-August-2026.docx"


def add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, value) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("Daily Activity Report (DAR)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{ORG}\n{PRODUCT} — {URL}\n")

    doc.add_paragraph()
    add_meta_table(
        doc,
        [
            ("Report date", REPORT_DATE),
            ("Prepared by", PREPARED_BY),
            ("Git branch (feature)", BRANCH),
            ("Base branch", BASE_BRANCH),
            ("Commit", "270d84e — Improve Aadhaar dual signature appearance and VSign live/UAT switch tooling"),
            ("Environment", "Local VSign live keys + production server pull started (esp.documantra.in)"),
            ("Report type", "Daily Activity Report (DAR)"),
        ],
    )
    doc.add_paragraph()

    sections = [
        (
            "1. Executive summary",
            [
                "Fixed Aadhaar eSign dual signature appearance: content-fit stamp, no wide white blank / double overlay after OTP, opaque handwritten bake on blue dual card.",
                "Polished pending (pre-OTP) dual UI: larger card, Times New Roman, wider tick; edit only via green pencil; Sign with Aadhaar CTA starts OTP even after handwritten save.",
                "Added VSign UAT↔live profile switch tooling (scripts + profiles); secrets remain gitignored; production callback URL supported for esp.documantra.in.",
                "Created and pushed feature branch feature/aadhaar-dual-appearance-vsign-live-21-8-26 (commit 270d84e); production VM fetched and checked out the branch.",
                "Did not commit live PFX, secrets/*.env, local Mongo data, utility logs, or JAR binaries.",
            ],
        ),
        (
            "2. Aadhaar dual signature appearance (UI + PDF)",
            [
                "Root causes addressed: SignPad white/transparent PNG baking to white in PDF; oversized React field wrap covering stamp; HTML dual overlay stacking on PDF bake after Aadhaar verify.",
                "Frontend: AadhaarSignatureAppearance, DocumentViewer, SignPad, vsignAppearance helpers — flatten/trim handwritten onto #E8F2FF; dual field width max-content; after verify show PDF only (no HTML overlay).",
                "Backend: vsignAppearanceEmbed.js — PDF dual paint, Times font, thicker tick, strip legacy wide SMask images; harden transparentize / isolate blend / blue top-right cover.",
                "UX: pencil-only edit for handwritten; box click does not open SignPad; Sign with Aadhaar always invokes doSign for OTP.",
            ],
        ),
        (
            "3. Local VSign live verification",
            [
                "Switched local profile UAT → live via switch-vsign-env.js (ASP IIPL001, esign.verasys.in, live PFX).",
                "Used Cloudflare quick tunnel for HTTPS VSign callback during local OTP testing; tunnels expire and must be refreshed.",
                "Services exercised locally: Vite 5173, e-sign 2103, pdf-service 2104, VSign utility 7078.",
                "DOCX upload path hardened (pdf-service mammoth fallback / timeouts) for Windows LibreOffice hangs.",
                "New envelopes required after live key switch; old UAT envelopes must not be reused for live OTP.",
            ],
        ),
        (
            "4. Git / branch publication",
            [
                "New branch: feature/aadhaar-dual-appearance-vsign-live-21-8-26 from prior recipient-portal work.",
                "Pushed to origin: https://github.com/mern-itio/esp.documantra.in (commit 270d84e).",
                "PR create URL available; gh CLI auth not configured on local workstation.",
                "Excluded from commit: .mongo-local, uploads/vSign/signCertificate.pfx changes, utility logs/JAR, Excel/XML dumps, config/vsign/secrets/*.env.",
            ],
        ),
        (
            "5. Production server status (21 Aug)",
            [
                "VM: DM-Indian-ubuntu-s-2vcpu-4gb-amd-blr1 — repo /root/Draft-and-Sign.",
                "Completed: git fetch; checkout feature/aadhaar-dual-appearance-vsign-live-21-8-26; branch up to date with origin (Already up to date after checkout = tip already at 270d84e).",
                "Restored incidental local edits: Frontend/package-lock.json, deploy/scripts/apply-nginx-vapt.sh.",
                "Left untracked (do not commit): Annexure-A16-UserConsents-Snapshot.json copies.",
                "Pending on server: switch-vsign-env live, rebuild/recreate e-sign-service + pdf-service, Frontend build → /var/www/draft-and-sign, smoke test with NEW envelope.",
            ],
        ),
        (
            "6. Commits pushed (21 August 2026)",
            [
                "270d84e — Improve Aadhaar dual signature appearance and VSign live/UAT switch tooling.",
            ],
        ),
        (
            "7. Key files touched",
            [
                "Frontend: DocumentViewer.tsx, AadhaarSignatureAppearance.tsx, SignPad.tsx, vsignAppearance.ts, vsignAadhaarStorage.ts, VSignAdminSettings.tsx.",
                "Backend e-sign: vSignController.js, aadhar.vsign.service.js, vsignAppearanceEmbed.js, vsignAssets.js, vsignConfigPolicy.js, callback middleware.",
                "Tooling: config/vsign/profiles/{live,uat}.json, switch-vsign-env.js, sync/update/verify scripts, deploy/scripts/deploy-vsign-live.sh.",
                "pdf-service: pdfController.js DOCX conversion resilience.",
            ],
        ),
        (
            "8. Testing & verification",
            [
                "Local: dual stamp sizing / no top-right white blank; handwritten opaque on blue; post-OTP PDF-only appearance.",
                "Local: pencil edits handwritten; Sign with Aadhaar starts OTP after handwritten save.",
                "Local live keys: VSign OTP path with tunnel callback (when tunnel active).",
                "Production: code present on disk after checkout; runtime still pending rebuild/restart at time of this DAR.",
            ],
        ),
        (
            "9. Production deploy steps (remaining)",
            [
                "cd /root/Draft-and-Sign && git log -1 --oneline  # expect 270d84e",
                "Confirm live PFX + config/vsign/secrets/live.env exist on server.",
                "cd Backend/services/e-sign-service && node scripts/switch-vsign-env.js live && node scripts/vsign-utility-props.js production && sudo systemctl restart vsign-utility || true",
                "docker compose -f docker-compose.prod.yml build e-sign-service pdf-service && up -d --force-recreate (or Backend/docker compose equivalent).",
                "Frontend: npm install --legacy-peer-deps && NODE_OPTIONS=--max-old-space-size=4096 npm run build && rsync dist/ → /var/www/draft-and-sign/",
                "Smoke: GET callback URL; create NEW envelope on https://esp.documantra.in; full Aadhaar eSign.",
            ],
        ),
        (
            "10. Pending / next steps",
            [
                "Complete production Docker rebuild + Frontend rsync; verify live Aadhaar eSign on esp.documantra.in.",
                "Open GitHub PR from feature branch into integration/main branch used for production.",
                "Confirm production callback remains https://esp.documantra.in/esign/api/e-sign/public/v-sign/response (no tunnel).",
                "Monitor first live signings for stamp appearance and OTP callback success.",
            ],
        ),
    ]

    for heading, bullets in sections:
        doc.add_heading(heading, 1)
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "This DAR is prepared for internal tracking and management review. "
        "Redact credentials and customer PII before external distribution."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
