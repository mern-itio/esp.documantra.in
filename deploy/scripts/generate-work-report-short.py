#!/usr/bin/env python3
"""Generate short work report for Documantra / ESP project."""
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "deploy" / "docs" / "Work-Report-Short.docx"

ORG = "ITIO Innovex Pvt Ltd"
DATE = datetime.now(timezone.utc).strftime("%d %B %Y")


def main() -> None:
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

    t = doc.add_heading("Work Report (Short)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{ORG}\nDocumantra / esp.documantra.in\nDate: {DATE}\n")

    sections = [
        (
            "1. Technical documentation",
            [
                "Created Documantra Technical Stack & Infrastructure reference (DOCX v1.1).",
                "Covers server, MongoDB, React frontend, Node.js microservices, nginx, integrations.",
                "Added detailed firewall & network security section (7-layer defence model).",
                "Path: deploy/docs/Documantra-Technical-Stack-and-Infrastructure.docx",
            ],
        ),
        (
            "2. ASP audit — database & security",
            [
                "Created Annexure A15: full MongoDB database/collection inventory (~75+ collections).",
                "Documented endpoint protection model (antivirus alternative for Linux cloud).",
                "Added production check scripts for DB and security verification.",
                "Paths: deploy/docs/asp-audit-annexures/Annexure-A15-*",
            ],
        ),
        (
            "3. Documantra marketing website (E:\\documantra)",
            [
                "Started local site (Vite + React + Supabase CMS) on localhost:8081.",
                "Fixed admin branding upload — Supabase RLS error (local login replaced with Supabase auth).",
                "Fixed logo/favicon not showing on live site — wired Header to Supabase branding storage.",
                "Added useBranding hook, BrandingHead (favicon), CMS sync after upload.",
            ],
        ),
        (
            "4. Production / ESP (esp.documantra.in)",
            [
                "Documented MongoDB: primary DB draftnsign, secondary support-db.",
                "Confirmed security model: Cloud Firewall, UFW, fail2ban, unattended-upgrades (no desktop AV).",
                "Provided PuTTY check script: deploy/scripts/check-production-db-and-security.sh",
            ],
        ),
        (
            "5. Deliverables summary",
            [
                "Documantra-Technical-Stack-and-Infrastructure.docx",
                "Annexure-A15-Database-Inventory-and-Endpoint-Protection.docx",
                "Annexure-A15-MongoDB-Collections.json",
                "check-production-db-and-security.sh + check-production-remote.bat",
                "supabase/ADMIN_SETUP.sql (documantra admin)",
            ],
        ),
        (
            "6. Pending / next steps",
            [
                "Run Supabase ADMIN_SETUP.sql if branding upload still needs admin role.",
                "Execute server DB check on production via PuTTY and attach output for ASP.",
                "Annexure A4 ISP updated (v1.1) with classification scheme and approval table — sign manually before auditor submission.",
                "Surepass e-Stamp Stamper V2 integration into main backend (sandbox proven locally).",
            ],
        ),
    ]

    for heading, bullets in sections:
        doc.add_heading(heading, 1)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph("Prepared for internal / ASP audit reference.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
