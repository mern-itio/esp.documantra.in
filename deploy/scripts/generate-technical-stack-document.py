#!/usr/bin/env python3
"""Generate Documantra Technical Stack & Infrastructure reference document."""
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "deploy" / "docs" / "Documantra-Technical-Stack-and-Infrastructure.docx"

ORG = "ITIO Innovex Pvt Ltd"
PRODUCT = "Documantra"
URL = "https://esp.documantra.in"
ADMIN_URL = f"{URL}/admin/"
ESIGN_PUBLIC_URL = "https://esign.documantra.in"
REGISTERED = (
    "C-32, Sector-14, Kaushambi, Ghaziabad, Uttar Pradesh 201010, India "
    "(CIN: U72300UP2015PTC070364)"
)
VERSION = "1.1"
DATE = datetime.now(timezone.utc).strftime("%d %B %Y")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("Technical Stack & Infrastructure Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{ORG}\n{PRODUCT} — Electronic Signature & Document Platform\n")
    run.bold = True
    sub.add_run(f"Production URL: {URL}\n")
    sub.add_run(f"Document version: {VERSION} | Generated: {DATE}")

    doc.add_paragraph()
    doc.add_paragraph(
        "यह दस्तावेज़ Documantra प्लेटफ़ॉर्म में उपयोग किए गए सर्वर, डेटाबेस, "
        "React frontend, Node.js backend और संबंधित तकनीकी विवरण का संपूर्ण संदर्भ है। "
        "This document is a consolidated reference for auditors, developers, and operations."
    )

    # 1. Executive summary
    doc.add_heading("1. Executive summary", 1)
    doc.add_paragraph(
        f"{PRODUCT} is an in-house developed Application Service Provider (ASP) platform for "
        "electronic document signing, envelope workflows, digital signatures (PKI), Aadhaar eSign "
        "(via V-Sign ESP), identity verification (Surepass DigiLocker), PDF tools, templates, "
        "subscriptions, and admin moderation."
    )
    add_bullets(
        doc,
        [
            f"Organization: {ORG}",
            f"Registered office: {REGISTERED}",
            f"Production application URL: {URL}",
            f"Admin panel URL: {ADMIN_URL}",
            f"Public signer SPA (optional host): {ESIGN_PUBLIC_URL}",
            "Hosting: DigitalOcean cloud — New York (NYC1) datacenter, United States",
            "Architecture: Microservices (Node.js) behind nginx reverse proxy with TLS termination",
        ],
    )

    # 2. Production server
    doc.add_heading("2. Production server & hosting", 1)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ("Cloud provider", "DigitalOcean"),
            ("Region / datacenter", "NYC1 (New York, United States)"),
            ("Server type", "Droplet — ubuntu-s-1vcpu-2gb-nyc1 (reference name: DandS)"),
            ("Operating system", "Ubuntu LTS"),
            ("Hostname / domain", "esp.documantra.in"),
            ("Public ports", "22 (SSH), 80 (HTTP redirect), 443 (HTTPS / HTTP2)"),
            ("TLS certificate", "Let's Encrypt (TLS 1.3)"),
            ("Reverse proxy", "nginx (host-level) — terminates TLS, routes to localhost services"),
            ("Application runtime", "Docker containers (microservices bound to 127.0.0.1)"),
            ("Static frontend (user SPA)", "/var/www/draft-and-sign"),
            ("Static frontend (admin SPA)", "/var/www/admin-esp"),
            ("Firewall", "DigitalOcean Cloud Firewall + UFW on host"),
            ("Patch management", "unattended-upgrades (automatic security updates)"),
            ("Intrusion prevention", "fail2ban (SSH and nginx)"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("2.1 Network flow", 2)
    doc.add_paragraph(
        "Internet → DigitalOcean Cloud Firewall (22/80/443) → Ubuntu host → nginx (443/HTTP2) "
        "→ Docker microservices on 127.0.0.1 ports OR static SPA files. "
        "MongoDB and internal APIs are not exposed to the public internet."
    )

    # 3. Frontend
    doc.add_heading("3. Frontend (React)", 1)
    add_table(
        doc,
        ["Technology", "Version / details"],
        [
            ("Framework", "React 18"),
            ("Language", "TypeScript"),
            ("Build tool", "Vite 7"),
            ("Routing", "React Router DOM 6"),
            ("UI libraries", "Material UI (MUI) 7, Radix UI, Tailwind CSS 4"),
            ("State management", "Zustand"),
            ("Forms / validation", "React Hook Form, Zod"),
            ("HTTP client", "Axios"),
            ("PDF viewing", "react-pdf, pdfjs-dist"),
            ("Document editor", "Fabric.js, CKEditor, TipTap"),
            ("Charts", "Chart.js, Recharts"),
            ("Real-time", "Socket.IO client"),
            ("Package name", "module-1 (Frontend/package.json)"),
            ("Build output", "Static SPA deployed to /var/www/draft-and-sign"),
            ("Admin frontend", "Separate repository — Draft-and-Sign-Admin-Frontend → /var/www/admin-esp"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("3.1 Key frontend routes", 2)
    add_bullets(
        doc,
        [
            "/ — User dashboard and application (SPA)",
            "/login, /auth-login — Authentication",
            "/admin/ — Admin panel (separate SPA build)",
            "Public signer flow — envelope signing without full login",
        ],
    )

    # 4. Backend
    doc.add_heading("4. Backend (Node.js microservices)", 1)
    doc.add_paragraph(
        "All backend services run on Node.js 20 (Docker base image: node:20). "
        "Each service is an independent Express.js application. Shared internal packages: "
        "auth-lib, validators, email-lib."
    )
    add_table(
        doc,
        ["Service", "Port", "nginx path prefix", "Primary responsibility"],
        [
            ("auth-service", "2101", "/auth/", "User/admin login, JWT, 2FA (TOTP), sessions"),
            ("document-service", "2102", "/document/", "Document upload, storage, sharing, analysis"),
            ("e-sign-service", "2103", "/esign/, /uploads/", "Envelopes, signing, PKI, audit trail, certificates"),
            ("pdf-service", "2104", "/pdf/", "PDF manipulation, form fill, extraction (Python helper)"),
            ("api-service", "2105", "/service/", "Public REST API analytics and endpoints"),
            ("template-service", "2106", "/template/", "Document templates"),
            ("support-service", "2107", "/support/", "Support tickets and agents"),
            ("ai-assistant-service", "2108", "/ai/", "AI assistant, RAG document embeddings"),
            ("subscription-service", "2110", "/subscription/", "Plans, billing, auth providers config"),
            ("organization-service", "2111", "/organization/", "Organization profiles and verification"),
            ("email-service", "2112", "(internal)", "Transactional email dispatch"),
            ("api-gateway", "2113", "/api/", "API gateway routing"),
            ("identity-service", "2114", "/identity/, /webhook/surepass-digilocker", "Aadhaar KYC via Surepass DigiLocker"),
            ("pdf-java-service", "2115", "(internal)", "Java-based PDF field operations"),
            ("admin-service", "3100", "/service/admin/", "Admin moderation APIs"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("4.1 Backend stack (common)", 2)
    add_table(
        doc,
        ["Component", "Details"],
        [
            ("Runtime", "Node.js 20"),
            ("Web framework", "Express.js 4.x"),
            ("ODM", "Mongoose 8.x"),
            ("Authentication", "JWT, bcrypt, httpOnly cookies, Google OAuth"),
            ("Security middleware", "Helmet, express-rate-limit, CORS, input validation"),
            ("PDF / crypto (e-sign)", "pdf-lib, node-signpdf, node-forge, jsrsasign, xml-crypto"),
            ("Email", "Nodemailer (SMTP)"),
            ("SMS / OTP", "Twilio (auth-service)"),
            ("Process manager (local dev)", "concurrently / nodemon"),
            ("Container orchestration (prod)", "Docker Compose"),
        ],
    )

    # 5. Database
    doc.add_heading("5. Database", 1)
    doc.add_paragraph(
        "Primary database: MongoDB. Connection string is configured per service via MONGO_URI "
        "environment variable (credentials stored in server .env — not committed to source control). "
        "MongoDB is not publicly accessible; it runs on the internal network / localhost."
    )
    add_table(
        doc,
        ["Item", "Details"],
        [
            ("Database engine", "MongoDB"),
            ("Access layer", "Mongoose ODM (Node.js)"),
            ("Connection config", "MONGO_URI in each microservice .env"),
            ("Known database names", "draftnsign (primary), support-db (support agents/tickets), document_management (document-service)"),
            ("Backup", "DigitalOcean snapshots / operational backup policy"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("5.1 Key MongoDB collections (by domain)", 2)
    add_table(
        doc,
        ["Domain / service", "Collections (Mongoose models)"],
        [
            ("auth-service", "Users, Admins, sessions, OTP logs"),
            ("e-sign-service", "Envelopes, Recipients, RecipientPermissions, Documents, DigitalSignatures, Certificates, AuditTrails, SignatureFields, SignatureTransactions, ActivityLogs, Notifications, Cycles, EnvelopeTypes, SelfSigners, OtpLogs"),
            ("document-service", "Documents, DocumentAnalysis"),
            ("organization-service", "Organizations"),
            ("subscription-service", "Subscriptions, plans, auth provider configs"),
            ("template-service", "Templates, template types"),
            ("support-service", "Tickets, SupportAgents (support-db)"),
            ("ai-assistant-service", "DocumentEmbeddings"),
            ("api-service", "ApiEndpointAnalytics"),
            ("identity-service", "Identity verification records (Surepass DigiLocker)"),
            ("admin-service", "Admin moderation data"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("5.2 Audit trail", 2)
    doc.add_paragraph(
        "Envelope signing actions are recorded in the AuditTrail collection (MongoDB) with timestamps "
        "and action details. Completion certificates and audit evidence are generated by e-sign-service."
    )

    # 6. nginx routing
    doc.add_heading("6. nginx reverse proxy routing", 1)
    doc.add_paragraph(
        "Production nginx configuration reference: deploy/nginx/esp.documantra.in.production.conf.example. "
        "All backend Docker ports are bound to 127.0.0.1 only; only nginx listens on 80/443."
    )
    add_bullets(
        doc,
        [
            "HTTP (port 80) → 301 redirect to HTTPS",
            "HTTPS (port 443, HTTP/2) → proxy to microservices or serve static SPA",
            "client_max_body_size: 100M (PDF uploads on /esign/)",
            "Security headers: HSTS, CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy",
            "VAPT hardening snippets: deploy/nginx/snippets/",
        ],
    )

    # 7. Integrations
    doc.add_heading("7. Third-party integrations", 1)
    add_table(
        doc,
        ["Integration", "Purpose", "Protocol"],
        [
            ("V-Sign ESP", "Aadhaar eSign (CCA standard eSign API v4.1)", "HTTPS + XML (EsignReq / EsignResp)"),
            ("Surepass DigiLocker", "Aadhaar identity verification (Via Link)", "HTTPS REST + webhook callback"),
            ("Surepass e-Stamp (Stamper V2)", "Electronic stamp paper generation (integration in progress)", "HTTPS REST"),
            ("Google OAuth", "Social login", "HTTPS OAuth 2.0"),
            ("Let's Encrypt", "TLS certificates", "ACME / certbot"),
            ("SMTP provider", "Transactional email", "SMTP (via email-service / Nodemailer)"),
            ("Twilio", "SMS OTP (optional)", "HTTPS API"),
        ],
    )

    # 8. Firewall & network security
    doc.add_heading("8. Firewall & network security", 1)
    doc.add_paragraph(
        "Documantra uses a defence-in-depth model: cloud-edge firewall, host firewall (UFW), "
        "localhost-only Docker bindings, nginx TLS termination, and application rate limits. "
        "MongoDB and all microservice APIs are never exposed directly to the public internet."
    )

    doc.add_heading("8.1 Defence layers (overview)", 2)
    add_table(
        doc,
        ["Layer", "Control", "Purpose"],
        [
            ("Layer 1 — Cloud edge", "DigitalOcean Cloud Firewall", "Block all inbound except 22, 80, 443 at datacenter edge"),
            ("Layer 2 — Host OS", "UFW (Uncomplicated Firewall)", "Deny all incoming by default; allow SSH/HTTP/HTTPS only"),
            ("Layer 3 — Docker bind", "127.0.0.1:PORT mapping", "Microservices reachable only from localhost (nginx proxy)"),
            ("Layer 4 — Reverse proxy", "nginx (443/HTTP2)", "TLS termination, routing, security headers, upload limits"),
            ("Layer 5 — Application", "express-rate-limit, Helmet, JWT", "Brute-force protection, secure headers, auth"),
            ("Layer 6 — Intrusion", "fail2ban", "SSH brute-force ban after repeated failures"),
            ("Layer 7 — Patching", "unattended-upgrades", "Automatic Ubuntu security patch installation"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("8.2 DigitalOcean Cloud Firewall", 2)
    doc.add_paragraph(
        "Configured in DigitalOcean console: Networking → Firewalls → attached to production droplet. "
        "Blocks ports at the cloud network edge even if host UFW is misconfigured."
    )
    add_table(
        doc,
        ["Rule type", "Protocol", "Port(s)", "Source", "Action"],
        [
            ("Inbound", "TCP", "22", "Trusted admin IPs (or 0.0.0.0/0 if required)", "ALLOW"),
            ("Inbound", "TCP", "80", "0.0.0.0/0, ::/0", "ALLOW"),
            ("Inbound", "TCP", "443", "0.0.0.0/0, ::/0", "ALLOW"),
            ("Inbound", "TCP", "2101–2115, 3100", "All", "DENY (implicit — not in allow list)"),
            ("Inbound", "TCP", "27017 (MongoDB)", "All", "DENY (not publicly exposed)"),
            ("Outbound", "All", "All", "Droplet", "ALLOW (default)"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("8.3 UFW host firewall", 2)
    doc.add_paragraph(
        "Script: deploy/scripts/harden-firewall.sh (also called by deploy/scripts/enable-vapt-production.sh). "
        "Run on droplet as root after Docker ports are bound to 127.0.0.1."
    )
    add_table(
        doc,
        ["UFW setting", "Value"],
        [
            ("Default incoming policy", "DENY"),
            ("Default outgoing policy", "ALLOW"),
            ("Rule: SSH", "ALLOW TCP 22"),
            ("Rule: HTTP", "ALLOW TCP 80 (redirect to HTTPS)"),
            ("Rule: HTTPS", "ALLOW TCP 443"),
            ("All other inbound", "BLOCKED"),
            ("Enable command", "ufw --force enable"),
            ("Verify", "ufw status verbose"),
        ],
    )
    doc.add_paragraph(
        "Optional hardening: restrict SSH to office IP only — "
        "ufw allow from YOUR.IP.ADDR.HERE to any port 22 proto tcp"
    )

    doc.add_paragraph()
    doc.add_heading("8.4 Docker port binding (localhost only)", 2)
    doc.add_paragraph(
        "Backend/docker-compose.yml binds every microservice to 127.0.0.1 only. "
        "Public access is exclusively via nginx on ports 80/443."
    )
    add_table(
        doc,
        ["Service", "Docker bind address", "Publicly accessible?"],
        [
            ("auth-service", "127.0.0.1:2101", "No — via nginx /auth/ only"),
            ("document-service", "127.0.0.1:2102", "No — via nginx /document/ only"),
            ("e-sign-service", "127.0.0.1:2103", "No — via nginx /esign/ only"),
            ("pdf-service", "127.0.0.1:2104", "No — via nginx /pdf/ only"),
            ("api-service", "127.0.0.1:2105", "No — via nginx /service/ only"),
            ("template-service", "127.0.0.1:2106", "No — via nginx /template/ only"),
            ("support-service", "127.0.0.1:2107", "No — via nginx /support/ only"),
            ("ai-assistant-service", "127.0.0.1:2108", "No — via nginx /ai/ only"),
            ("subscription-service", "127.0.0.1:2110", "No — via nginx /subscription/ only"),
            ("organization-service", "127.0.0.1:2111", "No — via nginx /organization/ only"),
            ("email-service", "127.0.0.1:2112", "No — internal only"),
            ("api-gateway", "127.0.0.1:2113", "No — via nginx /api/ only"),
            ("identity-service", "127.0.0.1:2114", "No — via nginx /identity/ only"),
            ("pdf-java-service", "127.0.0.1:2115", "No — internal only"),
            ("admin-service", "127.0.0.1:3100", "No — via nginx /service/admin/ only"),
            ("MongoDB", "localhost / internal", "No — never public"),
        ],
    )
    doc.add_paragraph(
        "Verification command on server: ss -tlnp | grep -E '210[0-9]|211[0-9]|3100' "
        "— all listeners must show 127.0.0.1, not 0.0.0.0."
    )

    doc.add_paragraph()
    doc.add_heading("8.5 nginx security (reverse proxy)", 2)
    add_table(
        doc,
        ["Control", "Details"],
        [
            ("TLS", "Let's Encrypt certificate; TLS 1.3; HTTP/2 enabled"),
            ("HTTP redirect", "Port 80 → 301 redirect to HTTPS"),
            ("server_tokens", "off — hides nginx version"),
            ("proxy_hide_header", "Strips Server and X-Powered-By from upstream responses"),
            ("HSTS", "Strict-Transport-Security: max-age=31536000; includeSubDomains"),
            ("CSP", "Content-Security-Policy on user/admin SPAs"),
            ("X-Frame-Options", "SAMEORIGIN (user SPA); DENY (admin SPA)"),
            ("X-Content-Type-Options", "nosniff"),
            ("Referrer-Policy", "strict-origin-when-cross-origin"),
            ("Permissions-Policy", "camera=(), microphone=(), geolocation=()"),
            ("OPTIONS blocking", "Returns 405 on OPTIONS for sensitive paths (VAPT)"),
            ("Upload limit", "client_max_body_size 100M on /esign/"),
            ("Custom error pages", "404.html, 502.html (no default nginx error pages)"),
            ("security.txt", "RFC 9116 at /security.txt"),
        ],
    )
    doc.add_paragraph("Config snippets: deploy/nginx/snippets/")
    add_bullets(
        doc,
        [
            "vapt-hide-server-header.conf — strip Server banner",
            "vapt-user-spa-headers.conf — user SPA security headers",
            "vapt-admin-spa-headers.conf — admin SPA headers (X-Frame-Options: DENY)",
            "vapt-esign-public-spa-headers.conf — public signer SPA headers",
        ],
    )

    doc.add_paragraph()
    doc.add_heading("8.6 fail2ban & patch management", 2)
    add_table(
        doc,
        ["Control", "Details"],
        [
            ("fail2ban", "Monitors SSH auth logs; bans IPs after repeated failed login attempts"),
            ("unattended-upgrades", "Automatic installation of Ubuntu security patches"),
            ("Docker isolation", "Services run in containers from trusted node:20 base images"),
            ("npm audit", "Run on release builds to check dependency vulnerabilities"),
            ("Backups", "DigitalOcean snapshots / operational backup policy"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("8.7 Application rate limits", 2)
    doc.add_paragraph("Implemented via express-rate-limit middleware in microservices.")
    add_table(
        doc,
        ["Service / endpoint", "Window", "Max requests", "Purpose"],
        [
            ("auth-service — login / 2FA / Google login", "1 minute", "3 per IP", "Brute-force protection"),
            ("auth-service — OTP send", "10 minutes", "3 per IP", "OTP abuse prevention"),
            ("auth-service — password reset", "15 minutes", "5 per IP", "Reset flooding prevention"),
            ("support-service — auth endpoints", "15 minutes", "5 per IP", "Auth brute-force"),
            ("support-service — messages", "1 minute", "30 per IP", "Spam prevention"),
            ("support-service — API", "15 minutes", "100 per IP", "General API abuse"),
            ("support-service — admin API", "15 minutes", "500 per IP", "Admin dashboard (higher limit)"),
            ("admin-service — mutations", "Per config", "Limited", "Admin write protection"),
            ("template-service — AI content", "Per config", "Limited", "AI endpoint abuse prevention"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("8.8 Firewall deployment scripts", 2)
    add_table(
        doc,
        ["Script", "Purpose"],
        [
            ("deploy/scripts/harden-firewall.sh", "Reset and configure UFW — allow 22/80/443 only"),
            ("deploy/scripts/enable-vapt-production.sh", "UFW + 2FA policy env + restart auth/e-sign services"),
            ("deploy/scripts/apply-nginx-vapt.sh", "Install VAPT nginx snippets and reload nginx"),
            ("deploy/scripts/deploy-vapt-live.sh", "Full VAPT production deploy (nginx + services)"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("8.9 Network diagram (text)", 2)
    doc.add_paragraph(
        "Internet\n"
        "  ↓\n"
        "DigitalOcean Cloud Firewall (allow: 22, 80, 443)\n"
        "  ↓\n"
        "Ubuntu Host — UFW (allow: 22, 80, 443; deny all other inbound)\n"
        "  ↓\n"
        "nginx :443 (TLS 1.3, HTTP/2, security headers)\n"
        "  ├─ /           → Static React SPA (/var/www/draft-and-sign)\n"
        "  ├─ /admin/     → Static Admin SPA (/var/www/admin-esp)\n"
        "  ├─ /auth/      → 127.0.0.1:2101 (auth-service)\n"
        "  ├─ /esign/     → 127.0.0.1:2103 (e-sign-service)\n"
        "  ├─ /document/  → 127.0.0.1:2102 (document-service)\n"
        "  ├─ /identity/  → 127.0.0.1:2114 (identity-service)\n"
        "  └─ …other paths → respective localhost ports\n"
        "  ↓\n"
        "MongoDB (internal — not public)\n"
        "External APIs: V-Sign ESP, Surepass (outbound HTTPS only)"
    )

    # 9. Application security
    doc.add_heading("9. Application security controls", 1)
    add_bullets(
        doc,
        [
            "TLS 1.3 on all public traffic; HTTP strictly redirected to HTTPS",
            "JWT + httpOnly secure cookies for session management",
            "Two-factor authentication (TOTP) for users and admins",
            "Password policy: min length, uppercase, lowercase, number, special character",
            "Session idle timeout: 8 hours; max concurrent sessions: 5 per user",
            "Login body encryption (RSA public key) on auth endpoints",
            "SSRF guard on logo URL ingestion; input validation on profile/org metadata",
            "HIDE_ERROR_DETAILS=true in production (no stack traces to clients)",
            "VAPT retest completed — 32 findings closed (deploy/docs/VAPT-RETEST-EVIDENCE.md)",
            "Evidence: Annexure A5 (hardening), A6 (server controls), A3 (TLS verification)",
        ],
    )

    # 10. Local development
    doc.add_heading("10. Local development environment", 1)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ("Backend root", "Backend/"),
            ("Frontend root", "Frontend/"),
            ("Start all services", "cd Backend && npm run dev (or npm run dev:core for essential services)"),
            ("Bootstrap env", "npm run bootstrap — copies Backend/.env to each service"),
            ("MongoDB (local)", "mongodb://localhost:27017 (MONGO_URI in .env)"),
            ("Frontend dev server", "cd Frontend && npm run dev (Vite, typically port 5173)"),
            ("Docker alternative", "docker compose up in Backend/ (all ports on 127.0.0.1)"),
            ("e-Stamp sandbox lab", "surepass-estamp-local/ — standalone Surepass Stamper V2 test app (port 5090)"),
        ],
    )

    # 11. Repository structure
    doc.add_heading("11. Repository structure", 1)
    add_bullets(
        doc,
        [
            "Frontend/ — React + Vite user SPA",
            "Backend/services/ — Node.js microservices (auth, e-sign, document, etc.)",
            "Backend/packages/ — Shared libraries (auth-lib, validators, email-lib)",
            "Backend/docker-compose.yml — Local/production Docker stack",
            "deploy/nginx/ — Production nginx configuration and VAPT snippets",
            "deploy/scripts/ — Deployment, audit document generators, firewall hardening",
            "deploy/docs/ — VAPT evidence, ASP audit annexures, this document",
            "surepass-estamp-local/ — Local Surepass e-Stamp integration lab",
        ],
    )

    # 12. Related audit documents
    doc.add_heading("12. Related audit & compliance documents", 1)
    add_table(
        doc,
        ["Document", "Path"],
        [
            ("ASP Preliminary Information (completed)", "deploy/docs/ASP-Audit-Preliminary-Information-Completed.docx"),
            ("Information Security Policy (A4)", "deploy/docs/asp-audit-annexures/Annexure-A4-Information-Security-Policy.docx"),
            ("ASP–ESP eSign Integration (A7)", "deploy/docs/asp-audit-annexures/Annexure-A7-ASP-ESP-eSign-Integration.docx"),
            ("Network diagram (A1)", "deploy/docs/asp-audit-annexures/Annexure-A1-Network-Diagram.png"),
            ("VAPT retest evidence", "deploy/docs/VAPT-RETEST-EVIDENCE.md"),
            ("System hardening report (A5)", "deploy/docs/asp-audit-annexures/Annexure-A5-System-Hardening-Report.docx"),
            ("Server security controls (A6)", "deploy/docs/asp-audit-annexures/Annexure-A6-Server-Security-Controls.png"),
            ("TLS verification (A3)", "deploy/docs/asp-audit-annexures/Annexure-A3-TLS-Verification.png"),
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("— End of document —").italic = True

    return doc


def main() -> None:
    doc = build()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
