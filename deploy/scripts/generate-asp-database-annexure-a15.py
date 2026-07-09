#!/usr/bin/env python3
"""Generate ASP annexure: full MongoDB inventory + endpoint protection (antivirus model)."""
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT_DOCX = ROOT / "deploy" / "docs" / "asp-audit-annexures" / "Annexure-A15-Database-Inventory-and-Endpoint-Protection.docx"
OUT_JSON = ROOT / "deploy" / "docs" / "asp-audit-annexures" / "Annexure-A15-MongoDB-Collections.json"

ORG = "ITIO Innovex Pvt Ltd"
PRODUCT = "Documantra (ASP)"
URL = "https://esp.documantra.in"
DATE = datetime.now(timezone.utc).strftime("%d %B %Y")

# MongoDB uses *collections* (not SQL tables). Mongoose model name -> typical collection name.
DATABASES = [
    {
        "name": "draftnsign",
        "role": "Primary application database (all microservices via MONGO_URI)",
        "services": [
            {
                "service": "auth-service",
                "collections": [
                    ("users", "User", "Registered accounts — email, 2FA, sessions, plan"),
                    ("adminusers", "AdminUser", "Admin panel accounts"),
                    ("sessionpolicyconfigs", "SessionPolicyConfig", "Session timeout / concurrent session policy"),
                    ("referrals", "Referral", "User referral records"),
                    ("referralprogramconfigs", "ReferralProgramConfig", "Referral program settings"),
                    ("rewards", "Reward", "Referral rewards"),
                    ("userconsents", "UserConsent", "User consent audit — ToS, privacy, marketing, e-sign disclosure with IP/timestamp"),
                ],
            },
            {
                "service": "e-sign-service",
                "collections": [
                    ("envelopes", "Envelope", "Signing envelopes — status, recipients, scheduling"),
                    ("recipients", "Recipient", "Envelope signers / viewers"),
                    ("recipientpermissions", "RecipientPermission", "Per-recipient signing permissions"),
                    ("documents", "Document", "Documents attached to envelopes (e-sign domain)"),
                    ("digitalsignatures", "DigitalSignature", "PKI digital signature records"),
                    ("certificates", "Certificate", "Signer certificates (ASP-issued RSA)"),
                    ("audittrails", "AuditTrail", "Immutable signing audit log — action, IP, timestamp, docHash"),
                    ("signaturefields", "SignatureFields", "Placed signature field coordinates"),
                    ("signaturetransactions", "SignatureTransaction", "Aadhaar eSign / ESP transaction refs"),
                    ("activitylogs", "ActivityLogs", "Envelope activity events"),
                    ("notifications", "Notification", "Signer notifications"),
                    ("envelopetypes", "EnvelopeType", "Envelope type definitions"),
                    ("cycles", "Cycle", "Recurring envelope cycles"),
                    ("otplogs", "OtpLog", "OTP send/verify audit"),
                    ("selfsigners", "SelfSigner", "Self-sign workflow data"),
                    ("userconsents", "UserConsent", "Consent records for signup and signer disclosure acceptance"),
                ],
            },
            {
                "service": "document-service",
                "collections": [
                    ("documents", "Document", "User document library (upload, share, version)"),
                    ("folders", "Folder", "Document folders"),
                    ("shareddocuments", "SharedDocument", "Shared document access links"),
                    ("versions", "Version", "Document version history"),
                    ("comments", "Comment", "Document comments"),
                    ("workflows", "Workflow", "Document approval workflows"),
                    ("documentanalyses", "DocumentAnalysis", "OCR / AI document analysis results"),
                ],
            },
            {
                "service": "organization-service",
                "collections": [
                    ("organizations", "Organization", "Organization profiles and verification"),
                    ("organizationusers", "OrganizationUser", "Org membership"),
                    ("organizationroles", "OrganizationRole", "Org roles"),
                    ("organizationpermissions", "OrganizationPermission", "Org permission matrix"),
                    ("orgfolders", "OrgFolder", "Organization document folders"),
                    ("orgfoldershares", "OrgFolderShare", "Shared org folders"),
                    ("folderenvelopes", "FolderEnvelope", "Envelopes linked to org folders"),
                ],
            },
            {
                "service": "subscription-service",
                "collections": [
                    ("subscriptions", "Subscription", "User/org subscription state"),
                    ("plantemplates", "PlanTemplate", "Subscription plan definitions"),
                    ("creditpackages", "CreditPackage", "Credit bundles"),
                    ("flexiblecreditpackages", "flexibleCreditPackage", "Flexible credit pricing"),
                    ("usagerecords", "UsageRecord", "Feature usage metering"),
                    ("invoices", "Invoice", "Billing invoices"),
                    ("authproviders", "AuthProvider", "DigiLocker / KYC provider config"),
                    ("toolsettings", "ToolSettings", "PDF tool feature flags"),
                    ("verificationcodes", "VerificationCode", "Phone/email verification codes"),
                ],
            },
            {
                "service": "template-service",
                "collections": [
                    ("templates", "template", "Document templates"),
                    ("templatetypes", "TemplateType", "Template categories"),
                    ("forms", "Form", "Form builder definitions"),
                    ("formfields", "FormFields", "Form field schemas"),
                    ("formsubmissions", "FormSubmission", "Submitted web forms"),
                    ("pendingdocuments", "PendingDocument", "AI template pending docs"),
                    ("aifeedbacks", "AIFeedback", "AI content feedback"),
                    ("feedbackcategories", "FeedbackCategory", "Feedback taxonomy"),
                ],
            },
            {
                "service": "pdf-service",
                "collections": [
                    ("pdfoperationtrackings", "PdfOperationTracking", "PDF tool operation logs"),
                    ("documenttrackings", "DocumentTracking", "PDF document tracking"),
                    ("commenteddocuments", "CommentedDocument", "PDF annotations"),
                    ("cloudserviceconnections", "CloudServiceConnection", "Google Drive / cloud OAuth"),
                    ("cloudfiles", "CloudFile", "Synced cloud files"),
                    ("oauthstates", "OAuthState", "OAuth CSRF state tokens"),
                    ("workflowtemplates", "WorkflowTemplate", "PDF workflow templates"),
                    ("workflowexecutions", "WorkflowExecution", "Workflow run history"),
                    ("guestusages", "GuestUsage", "Guest PDF tool limits"),
                    ("activesessions", "ActiveSession", "Active PDF editor sessions"),
                ],
            },
            {
                "service": "identity-service",
                "collections": [
                    ("identitysessions", "IdentitySession", "DigiLocker / KYC verification sessions"),
                    ("selfieverifications", "SelfieVerification", "Selfie auth verification records"),
                ],
            },
            {
                "service": "api-service",
                "collections": [
                    ("esignapikeys", "ESignApiKey", "Public API keys"),
                    ("apiendpointanalytics", "ApiEndpointAnalytics", "API usage analytics"),
                    ("analyticsdays", "AnalyticsDay", "Daily analytics aggregates"),
                    ("posts", "Post", "Community posts"),
                    ("supporttickets", "SupportTicket", "API-routed support tickets"),
                ],
            },
            {
                "service": "ai-assistant-service",
                "collections": [
                    ("documentembeddings", "DocumentEmbedding", "RAG vector embeddings"),
                    ("conversations", "Conversation", "AI chat sessions"),
                    ("useractions", "UserAction", "AI assistant user actions"),
                    ("learningpatterns", "LearningPattern", "AI learning patterns"),
                ],
            },
            {
                "service": "email-service",
                "collections": [
                    ("templates", "Template", "Email templates"),
                    ("smtpconfigurations", "SmtpConfiguration", "SMTP server config"),
                ],
            },
            {
                "service": "admin-service",
                "collections": [
                    ("pdftools", "PDFTool", "Admin PDF tool registry"),
                    ("pdftoolsettings", "PDFToolSettings", "PDF tool configuration"),
                    ("pdftoolactivations", "PDFToolActivation", "Tool activation per tenant"),
                ],
            },
        ],
    },
    {
        "name": "support-db",
        "role": "Support chat / ticketing (accessed via mongoose.connection.useDb)",
        "services": [
            {
                "service": "support-service",
                "collections": [
                    ("tickets", "Ticket", "Support tickets"),
                    ("messages", "Message", "Ticket chat messages"),
                    ("customers", "Customer", "Support customers"),
                    ("supportagents", "SupportAgent", "Support agents"),
                    ("typingindicators", "TypingIndicator", "Live chat typing state"),
                ],
            },
        ],
    },
]


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def build_docx() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    title = doc.add_heading("Database Inventory & Endpoint Protection", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{ORG}\n{PRODUCT}\n")
    r.bold = True
    sub.add_run(f"Production: {URL}\nAnnexure A15 | Generated: {DATE}")

    doc.add_paragraph()

    doc.add_heading("1. Purpose", 1)
    doc.add_paragraph(
        "This annexure provides ASP auditors with a complete inventory of production databases "
        "and MongoDB collections used by Documantra, and explains how server security / "
        "endpoint protection is implemented (including the antivirus protection model)."
    )

    doc.add_heading("2. Database technology", 1)
    add_table(
        doc,
        ["Item", "Detail"],
        [
            ("Database engine", "MongoDB (NoSQL document database)"),
            ("SQL tables", "Not applicable — MongoDB uses collections and documents (JSON/BSON)"),
            ("Access layer", "Mongoose ODM (Node.js) in each microservice"),
            ("Connection", "MONGO_URI environment variable per service (credentials in server .env)"),
            ("Public exposure", "No — MongoDB port 27017 not exposed to the internet"),
            ("Primary database", "draftnsign"),
            ("Secondary database", "support-db (support-service)"),
            ("Backup", "DigitalOcean snapshots / operational backup policy"),
            ("Host", "DigitalOcean Droplet — esp.documantra.in (Ubuntu, NYC1)"),
        ],
    )

    doc.add_heading("3. Complete database & collection inventory", 1)
    doc.add_paragraph(
        "Collection names below follow Mongoose default pluralization (lowercase). "
        "Verify live names on server with: mongosh draftnsign --eval \"db.getCollectionNames()\""
    )

    total = 0
    for db in DATABASES:
        doc.add_heading(f"3.{DATABASES.index(db) + 1} Database: {db['name']}", 2)
        doc.add_paragraph(db["role"])
        for svc in db["services"]:
            doc.add_heading(svc["service"], 3)
            rows = [[c[0], c[1], c[2]] for c in svc["collections"]]
            add_table(doc, ["Collection", "Model", "Description"], rows)
            total += len(rows)
            doc.add_paragraph()

    doc.add_paragraph(f"Total collections documented: {total}")

    doc.add_heading("4. Critical collections for ASP audit", 1)
    add_table(
        doc,
        ["Collection", "Audit relevance"],
        [
            ("users / adminusers", "Authentication, 2FA, session management"),
            ("envelopes / recipients", "e-Sign workflow and signer data"),
            ("audittrails", "Primary signing audit evidence (timestamp, IP, action, docHash)"),
            ("userconsents", "Explicit consent log — ToS, privacy, marketing, e-sign disclosure with IP/UA"),
            ("digitalsignatures / certificates", "PKI signing and certificate issuance"),
            ("signaturetransactions", "Aadhaar eSign ESP transaction linkage"),
            ("organizations", "Organization KYC and logo (SSRF-guarded)"),
            ("authproviders", "Surepass DigiLocker integration config"),
            ("identitysessions", "Aadhaar identity verification sessions"),
        ],
    )

    doc.add_heading("5. Sample audit trail document structure", 1)
    doc.add_paragraph("Collection: audittrails (e-sign-service)")
    add_table(
        doc,
        ["Field", "Type", "Description"],
        [
            ("envelopeId", "ObjectId", "Reference to envelope"),
            ("recipientId", "ObjectId", "Reference to recipient permission"),
            ("action", "String", "e.g. OTP_SENT, DOC_SIGNED, ENVELOPE_COMPLETED"),
            ("details", "Mixed (JSON)", "Action-specific metadata"),
            ("docHash", "String", "SHA-256 hash of document at action time"),
            ("ip", "String", "Client IP address"),
            ("userAgent", "String", "Browser user-agent"),
            ("timestamp", "Date", "UTC timestamp"),
        ],
    )
    doc.add_paragraph("Evidence: Annexure A11 (MongoDB screenshot + AuditTrail-Sample.json)")

    doc.add_heading("6. How to verify database on production (read-only)", 1)
    doc.add_paragraph("Run on server (PuTTY) as root:")
    code = doc.add_paragraph()
    code.add_run(
        "bash deploy/scripts/check-production-db-and-security.sh\n\n"
        "# Or manual:\n"
        "mongosh draftnsign --eval \"db.getCollectionNames()\"\n"
        "mongosh draftnsign --eval \"db.audittrails.countDocuments()\"\n"
        "ss -tlnp | grep 27017   # must show 127.0.0.1 only, not 0.0.0.0"
    ).font.size = Pt(9)

    doc.add_heading("7. Antivirus & endpoint protection — how it works", 1)
    doc.add_paragraph(
        "Documantra production runs on Ubuntu Linux (DigitalOcean cloud). "
        "A traditional desktop antivirus agent (e.g. Windows Defender, ClamAV daemon scanning all files) "
        "is NOT deployed on the server. This is intentional and aligns with common cloud VM security practice."
    )

    doc.add_heading("7.1 Protection model (compensating controls)", 2)
    add_table(
        doc,
        ["Control", "How it works", "What it protects against"],
        [
            (
                "DigitalOcean Cloud Firewall",
                "Inbound allow: TCP 22, 80, 443 only. All other ports denied at cloud edge.",
                "Port scanning, direct DB/API access from internet",
            ),
            (
                "UFW host firewall",
                "Default deny incoming; allow SSH/HTTP/HTTPS only.",
                "Misconfigured services exposed on host",
            ),
            (
                "Docker port binding",
                "Microservices bound to 127.0.0.1 (2101–2115, 3100). Only nginx is public.",
                "Direct access to Node.js APIs",
            ),
            (
                "fail2ban",
                "Monitors /var/log/auth.log; bans IPs after repeated SSH failures.",
                "SSH brute-force attacks",
            ),
            (
                "unattended-upgrades",
                "Automatic Ubuntu security patch installation.",
                "Known OS vulnerabilities (malware entry via exploits)",
            ),
            (
                "nginx TLS + headers",
                "TLS 1.3, HSTS, CSP, X-Frame-Options, rate limits on auth.",
                "MITM, XSS, clickjacking, credential stuffing",
            ),
            (
                "Application rate limits",
                "express-rate-limit on login, OTP, password reset.",
                "Brute-force on application layer",
            ),
            (
                "Docker image hygiene",
                "node:20 official base images; npm audit on releases.",
                "Vulnerable dependencies",
            ),
            (
                "MongoDB network isolation",
                "Not publicly reachable; credentials in .env only.",
                "Database exfiltration / ransomware",
            ),
        ],
    )

    doc.add_heading("7.2 Why no traditional antivirus agent?", 2)
    doc.add_paragraph(
        "Linux cloud application servers typically use a defence-in-depth model instead of periodic "
        "file-scanning antivirus: restricted network surface, automated patching, intrusion prevention, "
        "and container isolation. Desktop-style AV is more common on end-user workstations than on "
        "headless API servers. Documantra's ASP audit position (Annexure A6) documents this model."
    )

    doc.add_heading("7.3 How to verify security controls (auditor commands)", 2)
    p = doc.add_paragraph()
    p.add_run(
        "ufw status verbose\n"
        "systemctl status fail2ban\n"
        "systemctl status unattended-upgrades\n"
        "dpkg -l | grep -iE 'clamav|sophos|crowdstrike'   # typically empty (expected)\n"
        "ss -tlnp | grep -E '27017|210[0-9]'\n"
        "curl -sSI https://esp.documantra.in/ | grep -i strict-transport"
    ).font.size = Pt(9)

    doc.add_heading("7.4 Evidence references", 1)
    add_table(
        doc,
        ["Annexure", "Content"],
        [
            ("A5", "System Hardening & Infrastructure Security Report"),
            ("A6", "Server Security Controls (Linux cloud endpoint protection)"),
            ("A11", "MongoDB audittrails collection evidence"),
            ("A3", "TLS 1.3 verification"),
            ("A9", "VAPT retest — 32 findings closed"),
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph("— End of Annexure A15 —").alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def build_json() -> dict:
    import json

    collections = []
    for db in DATABASES:
        for svc in db["services"]:
            for coll, model, desc in svc["collections"]:
                collections.append(
                    {
                        "database": db["name"],
                        "service": svc["service"],
                        "collection": coll,
                        "mongooseModel": model,
                        "description": desc,
                    }
                )
    return {
        "generated": DATE,
        "organization": ORG,
        "product": PRODUCT,
        "productionUrl": URL,
        "databaseEngine": "MongoDB",
        "note": "MongoDB uses collections, not SQL tables.",
        "databases": [db["name"] for db in DATABASES],
        "totalCollections": len(collections),
        "collections": collections,
        "endpointProtection": {
            "traditionalAntivirusAgent": False,
            "model": "Linux cloud VM defence-in-depth",
            "controls": [
                "DigitalOcean Cloud Firewall",
                "UFW",
                "fail2ban",
                "unattended-upgrades",
                "Docker localhost binding",
                "nginx TLS and security headers",
                "application rate limits",
            ],
        },
    }


def main() -> None:
    import json

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    build_docx().save(OUT_DOCX)
    OUT_JSON.write_text(json.dumps(build_json(), indent=2), encoding="utf-8")
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_JSON}")


if __name__ == "__main__":
    main()
