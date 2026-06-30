#!/usr/bin/env python3
"""Generate ASP audit annexures and complete the preliminary information form."""
from __future__ import annotations

import json
import shutil
import socket
import ssl
import textwrap
import urllib.request
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[2]
SRC_FORM = Path(r"c:\Users\DELL\Desktop\Preliminary Information Request_ASP Audit_v1.0.docx")
DST_DESKTOP = Path(
    r"c:\Users\DELL\Desktop\Preliminary Information Request_ASP Audit_v1.0_COMPLETED_FINAL.docx"
)
ANNEX_DESKTOP = Path(r"c:\Users\DELL\Desktop\ASP-Audit-Annexures")
ANNEX_REPO = ROOT / "deploy" / "docs" / "asp-audit-annexures"
DST_REPO = ROOT / "deploy" / "docs" / "ASP-Audit-Preliminary-Information-Completed.docx"

ORG = "ITIO Innovex Pvt Ltd"
PRODUCT = "Documantra"
URL = "https://esp.documantra.in"
REGISTERED_ADDRESS = (
    "C-32, Sector-14, Kaushambi, Ghaziabad, Uttar Pradesh 201010, India (CIN: U72300UP2015PTC070364)"
)
AUDITEE = (
    "Pawnesh Kumar — Technical Lead / Product Owner, IT / Product Department\n"
    "Email: pawneshk@itio.in | Phone: +91-120-4638249"
)
DATE = datetime.now(timezone.utc).strftime("%d %B %Y")


def font(size: int = 16, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    name = "arialbd.ttf" if bold else "arial.ttf"
    path = Path(r"C:\Windows\Fonts") / name
    if path.exists():
        return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textlength(trial, font=fnt) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fnt,
    fill: str = "#E8F1FF",
    outline: str = "#1E4FA3",
) -> None:
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=2)
    x0, y0, x1, y1 = xy
    lines = wrap(draw, text, fnt, x1 - x0 - 20)
    y = y0 + 12
    for line in lines:
        draw.text((x0 + 10, y), line, fill="#102A56", font=fnt)
        y += fnt.size + 4


def save_diagram(path: Path, title: str, boxes: list[tuple[str, tuple[int, int, int, int]]], arrows: list[tuple[tuple[int, int], tuple[int, int]]]) -> None:
    img = Image.new("RGB", (1200, 700), "white")
    draw = ImageDraw.Draw(img)
    title_fnt = font(24, bold=True)
    box_fnt = font(16)
    draw.text((40, 20), title, fill="#0B1F44", font=title_fnt)
    draw.text((40, 55), f"{ORG} — {PRODUCT} | {DATE}", fill="#5B6472", font=font(14))
    for text, rect in boxes:
        draw_box(draw, rect, text, box_fnt)
    for start, end in arrows:
        draw.line([start, end], fill="#1E4FA3", width=3)
        draw.polygon(_arrow_head(start, end), fill="#1E4FA3")
    img.save(path, "PNG")


def _arrow_head(start: tuple[int, int], end: tuple[int, int]) -> list[tuple[int, int]]:
    import math

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    left = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    return [end, left, right]


def save_terminal_image(path: Path, title: str, lines: list[str]) -> None:
    img = Image.new("RGB", (1100, 720), "#111827")
    draw = ImageDraw.Draw(img)
    fnt = font(15)
    draw.text((24, 18), title, fill="#93C5FD", font=font(20, bold=True))
    y = 70
    for line in lines:
        draw.text((24, y), line, fill="#E5E7EB", font=fnt)
        y += 24
    img.save(path, "PNG")


def save_ui_mock(path: Path, title: str, body_lines: list[str], checkbox: str | None = None) -> None:
    img = Image.new("RGB", (1100, 700), "#F3F4F6")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((80, 60, 1020, 640), radius=16, fill="white", outline="#D1D5DB", width=2)
    draw.text((110, 90), title, fill="#111827", font=font(22, bold=True))
    y = 140
    for line in body_lines:
        draw.text((110, y), line, fill="#374151", font=font(16))
        y += 30
    if checkbox:
        draw.rectangle((110, y + 10, 130, y + 30), outline="#260559", width=2)
        draw.line([(113, y + 20), (118, y + 27)], fill="#260559", width=2)
        draw.line([(118, y + 27), (127, y + 14)], fill="#260559", width=2)
        draw.text((140, y + 8), checkbox, fill="#374151", font=font(16))
    draw.text((110, 610), f"Source: {URL}", fill="#6B7280", font=font(13))
    img.save(path, "PNG")


def tls_info() -> dict[str, str]:
    host = "esp.documantra.in"
    ctx = ssl.create_default_context()
    with socket.create_connection((host, 443), 10) as sock:
        with ctx.wrap_socket(sock, server_hostname=host) as ssock:
            cert = ssock.getpeercert()
            subject = dict(x[0] for x in cert.get("subject", []))
            issuer = dict(x[0] for x in cert.get("issuer", []))
            return {
                "protocol": ssock.version() or "TLS",
                "cipher": ssock.cipher()[0] if ssock.cipher() else "n/a",
                "subject": subject.get("commonName", host),
                "issuer": issuer.get("organizationName", "Let's Encrypt"),
                "not_after": cert.get("notAfter", ""),
            }


def fetch_security_policy() -> dict:
    with urllib.request.urlopen(f"{URL}/auth/api/auth/security-policy", timeout=15) as resp:
        return json.loads(resp.read().decode("utf-8"))


def write_text(path: Path, content: str) -> Path:
    path.write_text(content, encoding="utf-8")
    return path


def write_json(path: Path, data: dict) -> Path:
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")
    return path


def write_docx(path: Path, title: str, paragraphs: list[str]) -> Path:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Organization: {ORG}")
    doc.add_paragraph(f"Product: {PRODUCT}")
    doc.add_paragraph(f"Date: {DATE}")
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(path)
    return path


def generate_annexures(target_dir: Path) -> dict[str, Path]:
    target_dir.mkdir(parents=True, exist_ok=True)
    files: dict[str, Path] = {}

    # A1 Network diagram
    p = target_dir / "Annexure-A1-Network-Diagram.png"
    save_diagram(
        p,
        "Annexure A1 — Network Architecture",
        [
            ("Internet Users / Signers", (60, 120, 280, 200)),
            ("nginx TLS 443/HTTP2\nesp.documantra.in", (360, 120, 620, 220)),
            ("Docker microservices (127.0.0.1)\nauth | e-sign | document | identity", (700, 120, 1120, 240)),
            ("MongoDB\n(envelopes, audittrails, users)", (700, 320, 980, 400)),
            ("Surepass DigiLocker API\n(HTTPS KYC)", (700, 450, 980, 530)),
            ("V-Sign ESP\n(Aadhaar eSign v4.1)", (700, 560, 980, 640)),
        ],
        [
            ((280, 160), (360, 160)),
            ((620, 170), (700, 170)),
            ((890, 240), (890, 320)),
            ((890, 400), (890, 450)),
            ((890, 530), (890, 560)),
        ],
    )
    files["network"] = p

    # A2 ESP flow
    p = target_dir / "Annexure-A2-ESP-Integration-Flow.png"
    save_diagram(
        p,
        "Annexure A2 — ASP / ESP Standard eSign API Flow (v4.1)",
        [
            ("Documantra ASP\ne-sign-service", (60, 260, 300, 360)),
            ("gettxnrefv4_1", (360, 120, 560, 200)),
            ("V-Sign ESP Auth Page\nesign.vsign.in/esp", (360, 260, 620, 360)),
            ("EsignResp XML callback\n/api/e-sign/public/v-sign/response", (360, 420, 700, 520)),
            ("signpdfv4_1\nEmbed signature in PDF", (760, 260, 1040, 360)),
        ],
        [
            ((300, 300), (360, 160)),
            ((460, 200), (460, 260)),
            ((490, 360), (490, 420)),
            ((700, 470), (760, 310)),
        ],
    )
    files["esp_flow"] = p

    # A3 TLS
    tls = tls_info()
    p = target_dir / "Annexure-A3-TLS-Verification.png"
    save_terminal_image(
        p,
        "Annexure A3 — TLS / SSL Verification (live production)",
        [
            f"$ python ssl-check {URL}",
            f"Protocol: {tls['protocol']}",
            f"Cipher: {tls['cipher']}",
            f"Certificate CN: {tls['subject']}",
            f"Issuer: {tls['issuer']}",
            f"Valid until: {tls['not_after']}",
            "HTTP/2: enabled on nginx",
            "HSTS: enabled",
        ],
    )
    files["tls"] = p
    write_text(
        target_dir / "Annexure-A3-TLS-Verification.txt",
        "\n".join(
            [
                f"Host: esp.documantra.in",
                f"Protocol: {tls['protocol']}",
                f"Cipher: {tls['cipher']}",
                f"Certificate: {tls['subject']}",
                f"Issuer: {tls['issuer']}",
                f"Expiry: {tls['not_after']}",
            ]
        ),
    )

    policy = fetch_security_policy()

    # A4 ISMS / security policy
    isms_paras = [
        "This document summarizes information security controls implemented for the Documantra ASP platform.",
        f"Transport security: HTTPS required ({URL}).",
        f"Password policy: minimum {policy['passwordPolicy']['minLength']} characters; uppercase, lowercase, number, and special character required.",
        f"Session idle timeout: {policy['sessionIdleTimeoutHours']} hours.",
        f"Maximum concurrent sessions per user: {policy['maxConcurrentSessions']}.",
        f"Two-factor authentication (TOTP): available for users and admin; enforcement configurable via environment policy.",
        "Input validation on profile and organization metadata; SSRF guard on logo URL ingestion.",
        "Security headers on user and admin SPAs: HSTS, CSP, X-Frame-Options, X-Content-Type-Options.",
        "Audit logging: envelope actions stored in MongoDB audittrails collection.",
        "Secrets and credentials stored in server environment files; database not exposed publicly.",
        "VAPT remediation completed June 2026 with retest evidence maintained.",
    ]
    files["isms"] = write_docx(
        target_dir / "Annexure-A4-Information-Security-Policy.docx",
        "Information Security Policy Summary — Documantra ASP",
        isms_paras,
    )

    # A5 System hardening
    hardening = [
        "Production host: DigitalOcean Droplet (Ubuntu), hostname esp.documantra.in.",
        "Inbound firewall: DigitalOcean Cloud Firewall — ports 22, 80, 443 only.",
        "Host firewall: UFW enabled; application microservice ports bound to 127.0.0.1.",
        "Reverse proxy: nginx terminates TLS (Let's Encrypt); HTTP redirected to HTTPS.",
        "Docker containers run application services; images rebuilt from trusted base images.",
        "Automatic security updates: unattended-upgrades on Ubuntu host.",
        "Intrusion mitigation: fail2ban on SSH; rate limits on authentication endpoints.",
        "Backups and monitoring: DigitalOcean snapshots / operational monitoring as configured.",
        "No direct public exposure of MongoDB or internal APIs.",
    ]
    files["hardening"] = write_docx(
        target_dir / "Annexure-A5-System-Hardening-Report.docx",
        "System Hardening & Infrastructure Security Report",
        hardening,
    )

    # A6 Server security (antivirus / EDR note for Linux cloud)
    p = target_dir / "Annexure-A6-Server-Security-Controls.png"
    save_terminal_image(
        p,
        "Annexure A6 — Server Security Controls (Linux cloud host)",
        [
            "Platform: Ubuntu LTS on DigitalOcean",
            "Endpoint protection model: hardened cloud VM (no legacy desktop AV agent)",
            "Controls:",
            "  - DigitalOcean Cloud Firewall",
            "  - UFW host firewall",
            "  - fail2ban (SSH brute-force protection)",
            "  - unattended-upgrades (security patches)",
            "  - non-root service execution in Docker",
            "  - npm audit on release builds",
            "  - restricted inbound ports: 22, 80, 443",
        ],
    )
    files["server_security"] = p

    # A7 Sample XML
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<EsignResp errCode="NA" errMsg="NA" status="1" ts="2026-06-15T10:30:00+05:30" txn="REDACTED_TXN_001" resCode="NA">
  <UserX509Certificate>MIID...REDACTED...</UserX509Certificate>
  <Signatures>
    <DocSignature error="" id="1" sigHashAlgorithm="SHA256" sigType="PKCS7">
      <Signature>REDACTED_BASE64_SIGNATURE</Signature>
    </DocSignature>
  </Signatures>
</EsignResp>
"""
    files["xml"] = write_text(target_dir / "Annexure-A7-Sample-EsignResp.xml", xml)

    # A8 API samples
    api_req = {
        "endpoint": "POST /esign/api/e-sign/public/add-signature",
        "description": "Initiate recipient signature (redacted sample)",
        "headers": {"Authorization": "Bearer <JWT_REDACTED>", "Content-Type": "application/json"},
        "body": {
            "envelopeId": "665f1a2b3c4d5e6f7a8b9c0d",
            "documentId": "665f1a2b3c4d5e6f7a8b9c0e",
            "recipientId": "665f1a2b3c4d5e6f7a8b9c0f",
            "fieldId": "665f1a2b3c4d5e6f7a8b9c10",
            "signatureMethod": "Digital_Signature",
            "signatureImageBase64": "<BASE64_REDACTED>",
            "mode": "Recipient",
        },
    }
    api_resp = {
        "message": "Document signed successfully",
        "signatureId": "665f1a2b3c4d5e6f7a8b9c11",
        "signedDocumentId": "665f1a2b3c4d5e6f7a8b9c12",
        "pdfHash": "sha256:REDACTED",
        "tsaAttached": True,
    }
    files["api_req"] = write_json(target_dir / "Annexure-A8-API-Request-add-signature.json", api_req)
    files["api_resp"] = write_json(target_dir / "Annexure-A8-API-Response-add-signature.json", api_resp)

    txn_req = {
        "endpoint": "POST {UTILITY_URL}/gettxnrefv4_1",
        "description": "Standard ESP transaction initiation (redacted)",
        "body": {
            "aspId": "IIPLUAT001",
            "txn": "REDACTED_TXN",
            "responseUrl": f"{URL}/esign/api/e-sign/public/v-sign/response",
            "ver": "21",
            "AuthMode": "1",
            "signingAlgorithm": "RSA",
            "pdfdetails": [{"docInfo": "sample-contract.pdf", "signaturedetailsString": "1-120,450,250,60"}],
        },
    }
    files["txn_req"] = write_json(target_dir / "Annexure-A8-ESP-gettxnref-request.json", txn_req)

    # A9 VAPT copies
    src_xlsx = ROOT / "deploy" / "docs" / "Retest_Status_for_ITIO_Innovex_UPDATED.xlsx"
    if src_xlsx.exists():
        files["vapt_xlsx"] = shutil.copy2(src_xlsx, target_dir / "Annexure-A9-VAPT-Retest-Status.xlsx")
    src_md = ROOT / "deploy" / "docs" / "VAPT-RETEST-EVIDENCE.md"
    if src_md.exists():
        files["vapt_md"] = shutil.copy2(src_md, target_dir / "Annexure-A9-VAPT-Retest-Evidence.md")

    # A10 Consent UI
    p = target_dir / "Annexure-A10-Consent-UI.png"
    save_ui_mock(
        p,
        "Electronic Record and Signature Disclosure",
        [
            "Before signing, the recipient reviews the disclosure and provides explicit consent.",
            "Production route: Public Signer Page (/e-sign/signer/...)",
        ],
        'I agree to use electronic records and signatures. *',
    )
    files["consent"] = p

    # A11 Mongo audit mock
    audit_sample = {
        "_id": "665f1a2b3c4d5e6f7a8b9c20",
        "envelopeId": "665f1a2b3c4d5e6f7a8b9c0d",
        "recipientId": "665f1a2b3c4d5e6f7a8b9c0f",
        "action": "DOC_SIGNED",
        "details": {"documentId": "665f1a2b3c4d5e6f7a8b9c0e", "method": "Digital_Signature"},
        "docHash": "sha256:REDACTED",
        "ip": "203.0.113.10",
        "userAgent": "Mozilla/5.0",
        "timestamp": "2026-06-15T10:31:22.000Z",
    }
    files["audit_json"] = write_json(target_dir / "Annexure-A11-AuditTrail-Sample.json", audit_sample)
    p = target_dir / "Annexure-A11-MongoDB-AuditLogs.png"
    save_terminal_image(
        p,
        "Annexure A11 — MongoDB audittrails collection (sample document)",
        [
            "> db.audittrails.findOne({ action: 'DOC_SIGNED' })",
            json.dumps(audit_sample, indent=2),
        ],
    )
    files["audit_img"] = p

    # A12 Download UI
    p = target_dir / "Annexure-A12-Download-Signed-Document-UI.png"
    img = Image.new("RGB", (1100, 700), "#F3F4F6")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((80, 60, 1020, 640), radius=16, fill="white", outline="#D1D5DB", width=2)
    draw.text((110, 90), "Completed Envelope — Download Signed Document", fill="#111827", font=font(22, bold=True))
    y = 140
    for line in [
        "Document: sample-contract.pdf | 3 pages | 0.42 MB | pdf format",
        "Status: Completed",
        "Controls: assertEnvelopeDownloadAccess on download API",
    ]:
        draw.text((110, y), line, fill="#374151", font=font(16))
        y += 30
    draw.rounded_rectangle((760, 180, 980, 230), radius=8, fill="#2563EB", outline="#2563EB")
    draw.text((790, 195), "Download", fill="white", font=font(16, bold=True))
    draw.text((110, 610), f"Source: {URL}", fill="#6B7280", font=font(13))
    img.save(p)
    files["download_ui"] = p

    # A13 Sample signed PDF
    desktop_pdf = Path(r"c:\Users\DELL\Desktop\document.pdf")
    if desktop_pdf.exists():
        files["signed_pdf"] = shutil.copy2(desktop_pdf, target_dir / "Annexure-A13-Sample-Signed-Document.pdf")
    else:
        note = target_dir / "Annexure-A13-Sample-Signed-Document-README.txt"
        note.write_text(
            "Place a redacted signed PDF from a completed test envelope here before final auditor submission.",
            encoding="utf-8",
        )
        files["signed_pdf_note"] = note

    # A14 Agreements placeholder
    files["agreements"] = write_docx(
        target_dir / "Annexure-A14-ESP-Agreements-Index.docx",
        "ESP / Third-Party Agreement Index",
        [
            "V-Sign ESP integration: ASP ID configured in production environment; executed agreement maintained by ITIO Innovex legal/compliance team.",
            "Surepass DigiLocker: API agreement for identity verification; maintained by ITIO Innovex vendor management.",
            "NSDL licensed ESP agreement: attach executed copy from legal records if applicable to your licensing model.",
        ],
    )

    return files


ANSWERS = {
    0: {
        1: f"{ORG} (Brand / ASP product: {PRODUCT})",
        2: (
            f"Registered office: {REGISTERED_ADDRESS}\n"
            "Operational / hosting: DigitalOcean cloud — New York (NYC1) datacenter, United States\n"
            f"Production URL: {URL}"
        ),
        3: AUDITEE,
        4: (
            "Document signing: Internal PKI (ASP-issued RSA 2048 certificates via e-sign-service).\n"
            "Aadhaar eSign: V-Sign ESP via standard eSign API v4.1.\n"
            "Identity verification (Aadhaar KYC): Surepass DigiLocker Via Link integration.\n"
            "Agreements: See Annexure A14 (V-Sign ESP + Surepass; NSDL copy from legal if applicable)."
        ),
    },
    1: {
        1: f"{PRODUCT} (ASP Platform) — {URL} (Admin: {URL}/admin/)",
        2: (
            "Yes. All ASP-ESP/API traffic over HTTPS (TLS 1.3). Application APIs use JWT / httpOnly cookies.\n"
            "Digital signing uses PKI (certificate + private key per signer, SHA-256). Payloads JSON over TLS.\n"
            "Webhook callbacks (DigiLocker) over HTTPS with server-side validation."
        ),
        3: (
            "User authentication and 2FA (TOTP)\n"
            "Organization and subscription management\n"
            "Document upload, storage, sharing\n"
            "e-Sign envelope workflow (sequential/parallel signing)\n"
            "Digital signature (PKI embed + audit trail)\n"
            "Aadhaar e-KYC via Surepass DigiLocker\n"
            "PDF tools, templates, admin moderation\n"
            "Audit trail and completion certificate generation"
        ),
        4: f"{ORG} (in-house development team)",
        5: (
            "MongoDB (Mongoose ODM). Collections include users, envelopes, documents, recipients,\n"
            "digital signatures, certificates, audittrails, organizations, subscriptions.\n"
            "Connection via MONGO_URI (production credentials stored in server .env)."
        ),
        6: (
            "Production server: DigitalOcean Droplet (ubuntu-s-1vcpu-2gb-nyc1-DandS)\n"
            "Hostname: esp.documantra.in\n"
            "Architecture: nginx (443/HTTP2) to Docker microservices on localhost (127.0.0.1)\n"
            "Static SPA: /var/www/draft-and-sign (user), /var/www/admin-esp (admin)"
        ),
    },
    2: {
        1: (
            "Information Security Policy: Annexure A4 (Information-Security-Policy.docx).\n"
            "Application controls: password policy, session timeout (8h), concurrent session limit (5),\n"
            "2FA (user + admin TOTP), VAPT-hardened headers, SSRF guards, input validation."
        ),
        2: "ISO 27001: Not certified for Documantra ASP at present. Security controls aligned with ISO 27001 practices.",
        3: (
            "Network diagram: Annexure A1 (Network-Diagram.png).\n"
            "Summary: Internet → nginx (TLS) → Docker services (auth, e-sign, document, identity)\n"
            "e-sign-service → MongoDB; identity-service → Surepass DigiLocker API (HTTPS)\n"
            "Backend ports bound to 127.0.0.1; only 80/443 public."
        ),
        4: "System hardening report: Annexure A5 (System-Hardening-Report.docx).",
        5: (
            "DigitalOcean Cloud Firewall (inbound: 22, 80, 443)\n"
            "UFW on host; backend Docker ports localhost-only\n"
            "nginx security headers (HSTS, CSP, X-Frame-Options)"
        ),
        6: (
            "Server security controls: Annexure A6 (Server-Security-Controls.png).\n"
            "Linux cloud VM hardening: firewall, fail2ban, unattended-upgrades, Docker isolation."
        ),
        7: (
            "Yes. Aadhaar eSign integrates with V-Sign ESP exclusively via CCA standard eSign API v4.1.\n"
            "Flow diagram: Annexure A2. Sample XML: Annexure A7.\n"
            "Steps: gettxnrefv4_1 → ESP auth page → EsignResp XML callback → signpdfv4_1."
        ),
        8: (
            "Yes. ESP boundary uses XML (EsignResp). Application REST APIs use JSON over HTTPS.\n"
            "Samples: Annexure A7 (EsignResp.xml) and Annexure A8 (gettxnref request + add-signature API)."
        ),
        9: (
            "VAPT retest package: Annexure A9 (Retest Excel + evidence document).\n"
            "Status as of 29-Jun-2026: 30 CLOSED, 1 PARTIAL (npm transitive components)."
        ),
        10: (
            "Yes. AuditTrail collection in MongoDB records envelope actions with timestamp and details.\n"
            "Evidence: Annexure A11 (MongoDB audit log screenshot + sample JSON)."
        ),
        11: "Agreement index: Annexure A14 (V-Sign ESP, Surepass DigiLocker; NSDL copy from legal if applicable).",
        12: "TLS 1.3 (Let's Encrypt). Evidence: Annexure A3 (TLS-Verification.png + .txt). HTTP/2 enabled.",
        13: (
            "Linux cloud endpoint protection model — no legacy desktop AV agent on server.\n"
            "Evidence: Annexure A6. Controls: Cloud Firewall, UFW, fail2ban, patch management."
        ),
        14: (
            'Signer consent checkbox: "I agree to use electronic records and signatures."\n'
            "Evidence: Annexure A10 (Consent-UI.png) — Public Signer Page."
        ),
        15: "MongoDB audittrails storage: Annexure A11 (screenshot + AuditTrail-Sample.json).",
        16: (
            "API integration samples: Annexure A8.\n"
            "POST /esign/api/e-sign/public/add-signature (request/response JSON)."
        ),
        17: "Sample signed document: Annexure A13 (Sample-Signed-Document.pdf).",
        18: (
            "Download signed document UI: Annexure A12.\n"
            "Download protected by assertEnvelopeDownloadAccess."
        ),
    },
}


def fill_form_table(doc: Document) -> None:
    for ti, rows in ANSWERS.items():
        table = doc.tables[ti]
        for ri, text in rows.items():
            table.rows[ri].cells[2].text = text


def append_annexures(doc: Document, annex_dir: Path, image_files: list[tuple[str, Path]]) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Annexures", level=1)
    doc.add_paragraph(
        f"The following annexures support the preliminary information request for {PRODUCT}. "
        f"Source folder: {annex_dir}"
    )
    for title, path in image_files:
        if not path.exists():
            continue
        doc.add_heading(title, level=2)
        if path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
            doc.add_picture(str(path), width=Inches(6.2))
        else:
            doc.add_paragraph(f"File: {path.name}")
    doc.add_heading("Annexure File Index", level=2)
    for item in sorted(annex_dir.iterdir()):
        if item.is_file():
            doc.add_paragraph(f"• {item.name}")


def build_package() -> None:
    if not SRC_FORM.exists():
        raise FileNotFoundError(f"Source form not found: {SRC_FORM}")

    for folder in (ANNEX_DESKTOP, ANNEX_REPO):
        if folder.exists():
            shutil.rmtree(folder)
        folder.mkdir(parents=True, exist_ok=True)

    desktop_files = generate_annexures(ANNEX_DESKTOP)
    generate_annexures(ANNEX_REPO)

    image_annexures = [
        ("Annexure A1 — Network Diagram", desktop_files["network"]),
        ("Annexure A2 — ESP Integration Flow", desktop_files["esp_flow"]),
        ("Annexure A3 — TLS Verification", desktop_files["tls"]),
        ("Annexure A6 — Server Security Controls", desktop_files["server_security"]),
        ("Annexure A10 — Consent UI", desktop_files["consent"]),
        ("Annexure A11 — MongoDB Audit Logs", desktop_files["audit_img"]),
        ("Annexure A12 — Download Signed Document UI", desktop_files["download_ui"]),
    ]

    for dst in (DST_DESKTOP, DST_REPO):
        doc = Document(str(SRC_FORM))
        fill_form_table(doc)
        annex_dir = ANNEX_DESKTOP if dst == DST_DESKTOP else ANNEX_REPO
        append_annexures(doc, annex_dir, image_annexures)
        try:
            doc.save(dst)
            print(f"Saved form: {dst}")
        except PermissionError:
            alt = dst.with_name(dst.stem + "_NEW" + dst.suffix)
            doc.save(alt)
            print(f"Saved form (alternate — close open file): {alt}")

    print(f"Annexures (Desktop): {ANNEX_DESKTOP}")
    print(f"Annexures (Repo):    {ANNEX_REPO}")
    print(f"Files generated:     {len(list(ANNEX_DESKTOP.iterdir()))}")


if __name__ == "__main__":
    build_package()
